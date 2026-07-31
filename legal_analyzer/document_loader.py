from __future__ import annotations

import io
from email import policy
from email.parser import BytesParser
from pathlib import Path

import pytesseract
from docx import Document
from PIL import Image, ImageEnhance, ImageFilter, ImageOps

from .models import PageTrace
from .ocr_engine import OCRConfig, extract_pdf_pages
from .text_utils import clean_text, useful_characters


def _prep(image: Image.Image) -> Image.Image:
    image = ImageOps.grayscale(image)
    image = ImageOps.autocontrast(image)
    image = image.filter(ImageFilter.MedianFilter(size=3))
    return ImageEnhance.Contrast(image).enhance(1.5)


def _docx(name: str, content: bytes) -> list[PageTrace]:
    doc = Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                parts.append(" | ".join(values))
    text = clean_text("\n".join(parts))
    return [PageTrace(name, 1, text, "DOCX", useful_characters=useful_characters(text))]


def _txt(name: str, content: bytes) -> list[PageTrace]:
    text = ""
    for enc in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            text = content.decode(enc)
            break
        except UnicodeDecodeError:
            pass
    text = clean_text(text)
    return [PageTrace(name, 1, text, "TXT", useful_characters=useful_characters(text))]


def _image(name: str, content: bytes, lang: str) -> list[PageTrace]:
    image = _prep(Image.open(io.BytesIO(content)))
    data = pytesseract.image_to_data(
        image, lang=lang, config="--oem 3 --psm 6",
        output_type=pytesseract.Output.DICT,
    )
    words, confs = [], []
    for word, conf in zip(data.get("text", []), data.get("conf", [])):
        word = (word or "").strip()
        try:
            score = float(conf)
        except Exception:
            score = -1
        if word:
            words.append(word)
        if score >= 0:
            confs.append(score)
    text = clean_text(" ".join(words))
    avg = round(sum(confs) / len(confs), 2) if confs else 0.0
    return [PageTrace(
        name, 1, text, "OCR imagen",
        ocr_confidence=avg,
        useful_characters=useful_characters(text),
    )]


def _eml(name: str, content: bytes) -> list[PageTrace]:
    msg = BytesParser(policy=policy.default).parsebytes(content)
    parts = [
        f"De: {msg.get('From', '')}",
        f"Para: {msg.get('To', '')}",
        f"Fecha: {msg.get('Date', '')}",
        f"Asunto: {msg.get('Subject', '')}",
    ]
    if msg.is_multipart():
        for part in msg.walk():
            if part.get_content_type() == "text/plain":
                try:
                    parts.append(part.get_content())
                except Exception:
                    pass
    else:
        try:
            parts.append(msg.get_content())
        except Exception:
            pass
    text = clean_text("\n".join(parts))
    return [PageTrace(name, 1, text, "EML", useful_characters=useful_characters(text))]


def load_document(name: str, content: bytes, config: OCRConfig) -> list[PageTrace]:
    ext = Path(name).suffix.lower()
    if ext == ".pdf":
        return extract_pdf_pages(name, content, config)
    if ext == ".docx":
        return _docx(name, content)
    if ext == ".txt":
        return _txt(name, content)
    if ext in {".jpg", ".jpeg", ".png"}:
        return _image(name, content, config.language)
    if ext == ".eml":
        return _eml(name, content)
    raise ValueError(f"Formato no compatible: {ext}")
