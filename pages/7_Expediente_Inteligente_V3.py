from __future__ import annotations

import hashlib
import io

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from legal_analyzer.case_extractor import build_timeline, classify_document, extract_case_metadata
from legal_analyzer.document_loader import load_document
from legal_analyzer.models import PageTrace
from legal_analyzer.ocr_engine import OCRConfig, quality_label
from legal_analyzer.petition_comparator import compare_requests_with_answers, extract_requests


st.set_page_config(page_title="Expediente Inteligente V3", page_icon="🧠", layout="wide")
st.title("🧠 Expediente inteligente V3")
st.caption(
    "Carga PDF, Word, texto, imágenes y correos; extrae datos del proceso, "
    "construye una línea de tiempo y compara peticiones con respuestas."
)


def digest(content: bytes) -> str:
    return hashlib.sha256(content).hexdigest()


@st.cache_data(show_spinner=False, max_entries=50)
def cached_load(name, content_hash, content, enabled, min_chars, max_pages, dpi):
    del content_hash
    config = OCRConfig(
        enabled=enabled,
        min_useful_characters=min_chars,
        max_ocr_pages=max_pages,
        dpi=dpi,
    )
    return [x.to_dict() for x in load_document(name, content, config)]


def restore(data: dict) -> PageTrace:
    return PageTrace(**data)


def word_report(metadata, timeline, comparison) -> bytes:
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("INFORME DEL EXPEDIENTE INTELIGENTE V3")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_heading("1. Datos del proceso", level=1)
    for _, row in metadata.iterrows():
        p = doc.add_paragraph()
        r = p.add_run(f"{row['Campo']}: ")
        r.bold = True
        p.add_run(str(row["Valor"] or ""))

    doc.add_heading("2. Línea de tiempo", level=1)
    for _, row in timeline.iterrows():
        doc.add_paragraph(
            f"{row.get('Fecha principal', '')} — {row.get('Tipo', '')} — "
            f"{row.get('Documento', '')}"
        )

    doc.add_heading("3. Peticiones y respuestas", level=1)
    if comparison.empty:
        doc.add_paragraph("No se generó comparación.")
    else:
        for _, row in comparison.iterrows():
            doc.add_heading(f"Solicitud {row.get('N.º', '')}", level=2)
            doc.add_paragraph(str(row.get("Solicitud", "")))
            doc.add_paragraph(
                f"Respuesta: {row.get('Respuesta localizada', '')}"
            )
            doc.add_paragraph(
                f"Evaluación: {row.get('Evaluación revisada', '')}"
            )

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


with st.sidebar:
    st.header("Lectura y OCR")
    enabled = st.checkbox("Aplicar OCR cuando sea necesario", value=True)
    min_chars = st.slider("Mínimo de caracteres útiles", 20, 300, 80, 10)
    max_pages = st.slider("Máximo de páginas OCR por PDF", 5, 100, 40, 5)
    dpi = st.select_slider("Resolución OCR", [150, 200, 220, 250, 300], value=220)

files = st.file_uploader(
    "Sube las piezas del expediente",
    type=["pdf", "docx", "txt", "jpg", "jpeg", "png", "eml"],
    accept_multiple_files=True,
)

if not files:
    st.info("Carga las piezas del proceso para iniciar.")
    st.stop()

documents = {}
bar = st.progress(0)
msg = st.empty()

for i, uploaded in enumerate(files):
    msg.info(f"Procesando {uploaded.name}")
    content = uploaded.getvalue()
    try:
        raw = cached_load(
            uploaded.name, digest(content), content,
            enabled, min_chars, max_pages, dpi
        )
        documents[uploaded.name] = [restore(x) for x in raw]
    except Exception as error:
        st.error(f"No se pudo procesar {uploaded.name}: {error}")
        documents[uploaded.name] = []
    bar.progress((i + 1) / len(files))
msg.empty()

st.subheader("1. Datos generales")
metadata = pd.DataFrame([
    {"Campo": k, "Valor": v}
    for k, v in extract_case_metadata(documents).items()
])
metadata_edit = st.data_editor(
    metadata,
    use_container_width=True,
    hide_index=True,
    column_config={
        "Campo": st.column_config.TextColumn("Campo", disabled=True),
        "Valor": st.column_config.TextColumn("Valor", width="large"),
    },
    key="v3_metadata",
)

st.subheader("2. Línea de tiempo")
timeline = pd.DataFrame(build_timeline(documents))
if not timeline.empty:
    timeline["_orden"] = pd.to_datetime(timeline["Fecha principal"], errors="coerce")
    timeline = timeline.sort_values(["_orden", "Documento"], na_position="last").drop(columns="_orden")
timeline_edit = st.data_editor(
    timeline,
    use_container_width=True,
    hide_index=True,
    num_rows="dynamic",
    key="v3_timeline",
)

st.subheader("3. Calidad de lectura")
quality_rows = []
for name, pages in documents.items():
    kind, _ = classify_document(pages)
    for page in pages:
        quality_rows.append({
            "Documento": name,
            "Tipo": kind,
            "Página": page.page,
            "Método": page.extraction_method,
            "Confianza OCR": "" if page.ocr_confidence is None else f"{page.ocr_confidence:.1f}%",
            "Caracteres útiles": page.useful_characters,
            "Calidad": quality_label(page),
            "Advertencias": " | ".join(page.warnings),
            "Vista previa": page.text[:300],
        })
quality = pd.DataFrame(quality_rows)
st.dataframe(quality, use_container_width=True, hide_index=True)

st.subheader("4. Peticiones y respuestas")
names = list(documents)
candidates = [
    name for name, pages in documents.items()
    if classify_document(pages)[0] in {"Derecho de petición", "Acción de tutela"}
]
default_index = names.index(candidates[0]) if candidates else 0
petition_doc = st.selectbox(
    "Documento con solicitudes o pretensiones",
    names,
    index=default_index,
)
answer_options = [x for x in names if x != petition_doc]
answer_docs = st.multiselect(
    "Documentos de respuesta",
    answer_options,
    default=answer_options,
)

requests = extract_requests(documents.get(petition_doc, []))
answer_pages = [
    page for name in answer_docs for page in documents.get(name, [])
]
comparison = pd.DataFrame(compare_requests_with_answers(requests, answer_pages))

allowed = [
    "Respondida de fondo",
    "Respondida parcialmente",
    "Respuesta evasiva",
    "Sin respuesta",
    "No verificable",
    "Posible respuesta de fondo",
    "Respuesta parcial o relacionada",
    "Coincidencia débil; revisar",
    "Sin respuesta localizada",
]

if comparison.empty:
    st.warning("No se detectaron solicitudes con los patrones actuales.")
    comparison_edit = comparison
else:
    comparison_edit = st.data_editor(
        comparison,
        use_container_width=True,
        hide_index=True,
        column_config={
            "Evaluación revisada": st.column_config.SelectboxColumn(
                "Evaluación revisada",
                options=allowed,
                required=True,
            ),
            "Solicitud": st.column_config.TextColumn("Solicitud", width="large"),
            "Respuesta localizada": st.column_config.TextColumn(
                "Respuesta localizada",
                width="large",
            ),
            "Observaciones": st.column_config.TextColumn("Observaciones", width="large"),
        },
        key="v3_comparison",
    )

st.subheader("5. Exportación")
excel = io.BytesIO()
with pd.ExcelWriter(excel, engine="openpyxl") as writer:
    metadata_edit.to_excel(writer, sheet_name="Datos proceso", index=False)
    timeline_edit.to_excel(writer, sheet_name="Linea de tiempo", index=False)
    quality.to_excel(writer, sheet_name="Calidad lectura", index=False)
    comparison_edit.to_excel(writer, sheet_name="Peticiones respuestas", index=False)

c1, c2 = st.columns(2)
with c1:
    st.download_button(
        "Descargar Excel",
        excel.getvalue(),
        "expediente_inteligente_v3.xlsx",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True,
    )
with c2:
    st.download_button(
        "Descargar Word",
        word_report(metadata_edit, timeline_edit, comparison_edit),
        "informe_expediente_inteligente_v3.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )

st.warning(
    "La clasificación y la evaluación son preliminares. "
    "Corrige manualmente los campos antes de usar el informe."
)
