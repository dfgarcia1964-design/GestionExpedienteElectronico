from __future__ import annotations

import hashlib
import io
import re
from datetime import date, datetime, time, timedelta

import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Pt

from legal_analyzer.document_loader import load_document
from legal_analyzer.models import PageTrace
from legal_analyzer.ocr_engine import OCRConfig

from legal_ui.case_context import LOADED_FILES_KEY, apply_prefill
from legal_ui.colombia_calendar import calculate_deadline_simple as calculate_deadline
from legal_ui.colombia_calendar import merge_excluded_dates
from legal_ui.tool_bridge import render_active_case_banner, render_sync_terms_to_lexivox


st.set_page_config(
    page_title="Control de Términos",
    page_icon="⏳",
    layout="wide",
)

st.title("⏳ Control de Términos y Revisión de Documentos")
st.caption(
    "Sube un documento, detecta posibles términos y luego confirma el cómputo."
)

st.warning(
    "La detección es preliminar. Verifica siempre la norma aplicable, "
    "la fecha inicial, la notificación y si el término se cuenta en horas, "
    "días hábiles o calendario."
)


MONTHS = {
    "enero": 1,
    "febrero": 2,
    "marzo": 3,
    "abril": 4,
    "mayo": 5,
    "junio": 6,
    "julio": 7,
    "agosto": 8,
    "septiembre": 9,
    "octubre": 10,
    "noviembre": 11,
    "diciembre": 12,
}


def file_hash(content: bytes) -> str:
    return hashlib.sha256(content).hexdigest()


@st.cache_data(show_spinner=False, max_entries=40)
def cached_load(
    name: str,
    content_hash: str,
    content: bytes,
    enabled: bool,
    min_chars: int,
    max_pages: int,
    dpi: int,
) -> list[dict]:
    del content_hash

    config = OCRConfig(
        enabled=enabled,
        min_useful_characters=min_chars,
        max_ocr_pages=max_pages,
        dpi=dpi,
    )

    return [
        item.to_dict()
        for item in load_document(name, content, config)
    ]


def restore(data: dict) -> PageTrace:
    return PageTrace(**data)


def parse_date(text: str) -> date | None:
    numeric = re.search(
        r"\b([0-3]?\d)[/-]([01]?\d)[/-]((?:19|20)\d{2})\b",
        text,
    )

    if numeric:
        day, month, year = numeric.groups()

        try:
            return date(
                int(year),
                int(month),
                int(day),
            )
        except ValueError:
            pass

    textual = re.search(
        r"\b([0-3]?\d)\s+de\s+("
        + "|".join(MONTHS)
        + r")\s+de\s+((?:19|20)\d{2})\b",
        text.lower(),
    )

    if textual:
        day, month_name, year = textual.groups()

        try:
            return date(
                int(year),
                MONTHS[month_name],
                int(day),
            )
        except ValueError:
            pass

    return None


def detect_deadlines(pages: list[PageTrace]) -> list[dict]:
    rows: list[dict] = []

    term_patterns = [
        r"\b(?:dentro\s+de|en\s+el\s+t[ée]rmino\s+de|plazo\s+de)\s+"
        r"(\d{1,3})\s+(horas?|d[ií]as?)\b",
        r"\b(\d{1,3})\s+(horas?|d[ií]as?)\s+"
        r"(?:siguientes|contados|improrrogables|h[áa]biles|calendario)\b",
        r"\bpor\s+el\s+t[ée]rmino\s+de\s+"
        r"(\d{1,3})\s+(horas?|d[ií]as?)\b",
    ]

    action_terms = (
        "responder",
        "cumplir",
        "informar",
        "entregar",
        "remitir",
        "pronunciarse",
        "allegar",
        "aportar",
        "notificar",
        "realizar",
        "autorizar",
        "resolver",
        "presentar",
    )

    for page in pages:
        paragraphs = [
            fragment.strip()
            for fragment in re.split(
                r"(?<=[\.\;\:])\s+|\n+",
                page.text,
            )
            if fragment.strip()
        ]

        for fragment in paragraphs:
            normalized = fragment.lower()
            found = None

            for pattern in term_patterns:
                match = re.search(
                    pattern,
                    normalized,
                    flags=re.IGNORECASE,
                )

                if match:
                    found = match
                    break

            if not found:
                continue

            quantity = int(found.group(1))
            raw_unit = found.group(2).lower()
            unit = "Horas" if "hora" in raw_unit else "Días"

            rule = "Calendario"

            if unit == "Días":
                if "hábil" in normalized or "habil" in normalized:
                    rule = "Hábiles"
                elif "calendario" in normalized:
                    rule = "Calendario"
                else:
                    rule = "Por confirmar"

            action = next(
                (
                    term
                    for term in action_terms
                    if term in normalized
                ),
                "Actuación por identificar",
            )

            detected_date = parse_date(fragment)

            rows.append(
                {
                    "Documento": page.document,
                    "Página": page.page,
                    "Fragmento completo": fragment,
                    "Cantidad detectada": quantity,
                    "Unidad detectada": unit,
                    "Regla detectada": rule,
                    "Acción detectada": action,
                    "Fecha mencionada": detected_date,
                    "Fecha inicial confirmada": None,
                    "Hora inicial": time(8, 0),
                    "Aplicar al control": True,
                    "Observaciones": "",
                }
            )

    unique: list[dict] = []
    seen: set[tuple] = set()

    for row in rows:
        key = (
            row["Documento"],
            row["Página"],
            row["Fragmento completo"],
        )

        if key not in seen:
            seen.add(key)
            unique.append(row)

    return unique


def status_for(deadline: datetime) -> tuple[str, str, float]:
    hours = (deadline - datetime.now()).total_seconds() / 3600

    if hours < 0:
        return "Rojo", "Vencido", hours
    if hours <= 24:
        return "Rojo", "Urgente", hours
    if hours <= 72:
        return "Amarillo", "Próximo", hours

    return "Verde", "En plazo", hours


def remaining_text(hours: float) -> str:
    if hours < 0:
        return f"Vencido hace {abs(hours):.1f} horas"

    days = int(hours // 24)
    extra = hours % 24

    if days:
        return f"{days} día(s) y {extra:.1f} hora(s)"

    return f"{extra:.1f} hora(s)"


def suggested_action(color: str, action: str) -> str:
    action = action or "la actuación pendiente"

    if color == "Rojo":
        return (
            f"Verificar inmediatamente el cómputo y preparar {action}, "
            "con anexos y prueba de envío."
        )

    if color == "Amarillo":
        return (
            f"Dejar listo el borrador de {action} y confirmar hoy la fecha exacta."
        )

    return (
        f"Programar la preparación de {action} y conservar el soporte "
        "de la fecha inicial."
    )


if "terms" not in st.session_state:
    st.session_state.terms = []


with st.sidebar:
    st.header("OCR")
    ocr_enabled = st.checkbox(
        "Aplicar OCR",
        value=True,
    )
    min_chars = st.slider(
        "Mínimo de caracteres útiles",
        20,
        300,
        80,
        10,
    )
    max_pages = st.slider(
        "Máximo de páginas OCR",
        5,
        100,
        40,
        5,
    )
    dpi = st.select_slider(
        "Resolución OCR",
        [150, 200, 220, 250, 300],
        value=220,
    )
    usar_festivos_co = st.checkbox(
        "Excluir festivos Colombia (días hábiles)",
        value=True,
    )


render_active_case_banner()
apply_prefill(
    {
        "manual_expediente": "caso_nombre",
        "manual_radicado": "radicado",
    }
)

st.subheader("1. Subir documento para revisar términos")

uploaded = st.file_uploader(
    "Sube el auto, fallo, requerimiento, respuesta o constancia",
    type=["pdf", "docx", "txt", "jpg", "jpeg", "png", "eml"],
    accept_multiple_files=False,
)

if not uploaded:
    loaded = st.session_state.get(LOADED_FILES_KEY) or []
    if loaded:
        uploaded = loaded[0]

detected_df = pd.DataFrame()

if uploaded:
    content = uploaded.getvalue()

    with st.spinner("Leyendo documento y buscando posibles términos..."):
        raw = cached_load(
            uploaded.name,
            file_hash(content),
            content,
            ocr_enabled,
            min_chars,
            max_pages,
            dpi,
        )

        pages = [
            restore(item)
            for item in raw
        ]

        detected = detect_deadlines(pages)
        detected_df = pd.DataFrame(detected)

    if detected_df.empty:
        st.warning(
            "No se detectaron expresiones claras de plazo. "
            "Puedes registrar el término manualmente más abajo."
        )
    else:
        st.success(
            f"Se detectaron {len(detected_df)} posible(s) término(s)."
        )

        edited_detected = st.data_editor(
            detected_df,
            use_container_width=True,
            hide_index=True,
            column_config={
                "Fragmento completo": st.column_config.TextColumn(
                    "Fragmento completo",
                    width="large",
                ),
                "Unidad detectada": st.column_config.SelectboxColumn(
                    "Unidad",
                    options=["Horas", "Días"],
                ),
                "Regla detectada": st.column_config.SelectboxColumn(
                    "Regla",
                    options=[
                        "Hábiles",
                        "Calendario",
                        "Por confirmar",
                    ],
                ),
                "Fecha inicial confirmada": st.column_config.DateColumn(
                    "Fecha inicial confirmada",
                    format="YYYY-MM-DD",
                ),
                "Hora inicial": st.column_config.TimeColumn(
                    "Hora inicial",
                    format="HH:mm",
                ),
                "Aplicar al control": st.column_config.CheckboxColumn(
                    "Aplicar",
                ),
                "Observaciones": st.column_config.TextColumn(
                    "Observaciones",
                    width="large",
                ),
            },
            key="detected_terms_editor",
        )

        excluded_text = st.text_area(
            "Fechas excluidas adicionales",
            placeholder="Una por línea: AAAA-MM-DD",
            height=90,
            key="detected_excluded_dates",
        )

        if st.button(
            "Agregar términos confirmados a la agenda",
            type="primary",
            use_container_width=True,
        ):
            st.session_state["usar_festivos_co"] = usar_festivos_co
            excluded, invalid = merge_excluded_dates(
                excluded_text,
                include_colombia=usar_festivos_co,
            )

            if invalid:
                st.error(
                    "Fechas inválidas: " + ", ".join(invalid)
                )
            else:
                added = 0

                for _, row in edited_detected.iterrows():
                    if not bool(row.get("Aplicar al control")):
                        continue

                    confirmed_date = row.get(
                        "Fecha inicial confirmada"
                    )

                    if pd.isna(confirmed_date):
                        st.warning(
                            "Debes confirmar la fecha inicial de cada término seleccionado."
                        )
                        continue

                    rule = str(row.get("Regla detectada", ""))

                    if rule == "Por confirmar":
                        st.warning(
                            "Debes confirmar si el término es hábil o calendario."
                        )
                        continue

                    start_time = row.get("Hora inicial")

                    if pd.isna(start_time):
                        start_time = time(8, 0)

                    if isinstance(confirmed_date, pd.Timestamp):
                        confirmed_date = confirmed_date.date()

                    start = datetime.combine(
                        confirmed_date,
                        start_time,
                    )

                    quantity = int(
                        row.get("Cantidad detectada", 0)
                    )
                    unit = str(
                        row.get("Unidad detectada", "Días")
                    )

                    deadline = calculate_deadline(
                        start,
                        quantity,
                        unit,
                        rule,
                        excluded,
                    )

                    color, state, hours = status_for(
                        deadline
                    )

                    st.session_state.terms.append(
                        {
                            "Expediente": uploaded.name,
                            "Radicado": "",
                            "Actuación": row.get(
                                "Acción detectada",
                                "",
                            ),
                            "Hecho inicial": row.get(
                                "Fragmento completo",
                                "",
                            ),
                            "Documento fuente": row.get(
                                "Documento",
                                "",
                            ),
                            "Página fuente": row.get(
                                "Página",
                                "",
                            ),
                            "Fecha inicial": start,
                            "Término": (
                                f"{quantity} "
                                f"{unit.lower()}"
                            ),
                            "Regla": rule,
                            "Vencimiento": deadline,
                            "Semáforo": color,
                            "Estado": state,
                            "Tiempo restante": remaining_text(
                                hours
                            ),
                            "Qué hacer": suggested_action(
                                color,
                                str(
                                    row.get(
                                        "Acción detectada",
                                        "",
                                    )
                                ),
                            ),
                            "Soporte": (
                                f"{uploaded.name}, "
                                f"página {row.get('Página', '')}"
                            ),
                            "Revisión humana": row.get(
                                "Observaciones",
                                "",
                            ),
                        }
                    )

                    added += 1

                if added:
                    st.success(
                        f"Se agregaron {added} término(s) a la agenda."
                    )


st.subheader("2. Registrar término manualmente")

with st.form("manual_term"):
    c1, c2, c3 = st.columns(3)

    with c1:
        expediente = st.text_input("Expediente", key="manual_expediente")
        radicado = st.text_input("Radicado", key="manual_radicado")
        actuacion = st.text_input(
            "Próxima actuación",
            placeholder="Ejemplo: incidente de desacato",
        )

    with c2:
        hecho = st.text_input(
            "Hecho que inicia el término",
            placeholder="Ejemplo: notificación del auto",
        )
        fecha = st.date_input(
            "Fecha inicial",
            value=date.today(),
        )
        hora = st.time_input(
            "Hora inicial",
            value=time(8, 0),
        )

    with c3:
        cantidad = st.number_input(
            "Cantidad",
            min_value=0,
            value=1,
            step=1,
        )
        unidad = st.selectbox(
            "Unidad",
            ["Días", "Horas"],
        )
        regla = st.selectbox(
            "Regla",
            ["Hábiles", "Calendario"],
            disabled=unidad == "Horas",
        )

    soporte = st.text_area(
        "Soporte de la fecha inicial",
        placeholder="Documento, página, correo o constancia.",
        height=90,
    )

    guardar = st.form_submit_button(
        "Calcular y guardar manualmente",
        use_container_width=True,
    )


if guardar:
    start = datetime.combine(fecha, hora)
    excluded_manual, _ = merge_excluded_dates(
        "",
        include_colombia=st.session_state.get("usar_festivos_co", True),
    )
    deadline = calculate_deadline(
        start,
        int(cantidad),
        unidad,
        regla,
        excluded_manual,
    )
    color, state, hours = status_for(deadline)

    st.session_state.terms.append(
        {
            "Expediente": expediente,
            "Radicado": radicado,
            "Actuación": actuacion,
            "Hecho inicial": hecho,
            "Documento fuente": "",
            "Página fuente": "",
            "Fecha inicial": start,
            "Término": f"{int(cantidad)} {unidad.lower()}",
            "Regla": (
                regla
                if unidad == "Días"
                else "Horas continuas"
            ),
            "Vencimiento": deadline,
            "Semáforo": color,
            "Estado": state,
            "Tiempo restante": remaining_text(hours),
            "Qué hacer": suggested_action(
                color,
                actuacion,
            ),
            "Soporte": soporte,
            "Revisión humana": "",
        }
    )

    st.success(
        f"Vencimiento calculado: {deadline:%Y-%m-%d %H:%M}"
    )


st.subheader("3. Agenda procesal")

if not st.session_state.terms:
    st.info("Todavía no hay términos registrados.")
    st.stop()


table = pd.DataFrame(
    st.session_state.terms
)

priority = {
    "Rojo": 0,
    "Amarillo": 1,
    "Verde": 2,
}

table["_orden"] = (
    table["Semáforo"]
    .map(priority)
    .fillna(9)
)

table = (
    table.sort_values(
        ["_orden", "Vencimiento"],
        ascending=[True, True],
    )
    .drop(columns="_orden")
)

c1, c2, c3, c4 = st.columns(4)
c1.metric("Total", len(table))
c2.metric(
    "🔴 Urgentes",
    int(
        (table["Semáforo"] == "Rojo").sum()
    ),
)
c3.metric(
    "🟡 Próximos",
    int(
        (table["Semáforo"] == "Amarillo").sum()
    ),
)
c4.metric(
    "🟢 En plazo",
    int(
        (table["Semáforo"] == "Verde").sum()
    ),
)

edited = st.data_editor(
    table,
    use_container_width=True,
    hide_index=True,
    num_rows="dynamic",
    column_config={
        "Fecha inicial": st.column_config.DatetimeColumn(
            "Fecha inicial",
            format="YYYY-MM-DD HH:mm",
        ),
        "Vencimiento": st.column_config.DatetimeColumn(
            "Vencimiento",
            format="YYYY-MM-DD HH:mm",
        ),
        "Semáforo": st.column_config.SelectboxColumn(
            "Semáforo",
            options=[
                "Rojo",
                "Amarillo",
                "Verde",
            ],
        ),
        "Revisión humana": st.column_config.TextColumn(
            "Revisión humana",
            width="large",
        ),
    },
    key="terms_editor",
)

sync_rows = edited.to_dict(orient="records")
render_sync_terms_to_lexivox(
    sync_rows,
    source="control",
    key="sync_control_terms_agenda",
)

st.subheader("4. Qué hacer primero")

urgent = edited[
    edited["Semáforo"].isin(
        ["Rojo", "Amarillo"]
    )
]

if urgent.empty:
    st.success(
        "No hay vencimientos urgentes."
    )
else:
    for _, row in urgent.iterrows():
        icon = (
            "🔴"
            if row["Semáforo"] == "Rojo"
            else "🟡"
        )

        st.warning(
            f"{icon} **{row['Expediente'] or 'Expediente sin nombre'}**\n\n"
            f"**Actuación:** {row['Actuación'] or 'Sin definir'}\n\n"
            f"**Documento:** {row.get('Documento fuente', '')} "
            f"— página {row.get('Página fuente', '')}\n\n"
            f"**Vencimiento:** {row['Vencimiento']}\n\n"
            f"**Acción:** {row['Qué hacer']}"
        )


st.subheader("5. Exportar")

excel = io.BytesIO()

with pd.ExcelWriter(
    excel,
    engine="openpyxl",
) as writer:
    edited.to_excel(
        writer,
        sheet_name="Control de términos",
        index=False,
    )


def create_word(dataframe: pd.DataFrame) -> bytes:
    document = Document()
    document.styles["Normal"].font.name = "Arial"
    document.styles["Normal"].font.size = Pt(10)
    document.add_heading(
        "CONTROL DE TÉRMINOS Y PRÓXIMAS ACTUACIONES",
        level=0,
    )

    for index, row in dataframe.iterrows():
        document.add_heading(
            f"{index + 1}. {row.get('Expediente', '')}",
            level=1,
        )

        for label, field in (
            ("Radicado", "Radicado"),
            ("Actuación", "Actuación"),
            ("Documento fuente", "Documento fuente"),
            ("Página fuente", "Página fuente"),
            ("Fecha inicial", "Fecha inicial"),
            ("Término", "Término"),
            ("Regla", "Regla"),
            ("Vencimiento", "Vencimiento"),
            ("Estado", "Estado"),
            ("Semáforo", "Semáforo"),
            ("Qué hacer", "Qué hacer"),
            ("Soporte", "Soporte"),
            ("Revisión humana", "Revisión humana"),
        ):
            document.add_paragraph(
                f"{label}: {row.get(field, '')}"
            )

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


d1, d2 = st.columns(2)

with d1:
    st.download_button(
        "Descargar agenda en Excel",
        data=excel.getvalue(),
        file_name="control_terminos_documentos.xlsx",
        mime=(
            "application/vnd.openxmlformats-officedocument."
            "spreadsheetml.sheet"
        ),
        use_container_width=True,
    )

with d2:
    st.download_button(
        "Descargar agenda en Word",
        data=create_word(edited),
        file_name="control_terminos_documentos.docx",
        mime=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
        use_container_width=True,
    )


if st.button(
    "Eliminar todos los términos",
):
    st.session_state.terms = []
    st.rerun()
