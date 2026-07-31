from __future__ import annotations

import hashlib
import io

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from legal_analyzer.case_extractor import (
    build_timeline,
    classify_document,
    extract_case_metadata,
)
from legal_analyzer.document_loader import load_document
from legal_analyzer.models import PageTrace
from legal_analyzer.ocr_engine import OCRConfig, quality_label
from legal_analyzer.question_engine import (
    answer_question,
    build_fragment_index,
    suggested_questions,
)


st.set_page_config(
    page_title="Sala de Estrategia Jurídica",
    page_icon="🧭",
    layout="wide",
)

st.title("🧭 Sala de Estrategia Jurídica")
st.caption(
    "Pregunta al expediente y recibe respuestas con documento, página, "
    "fragmento y nivel de confianza."
)


def sha256(content: bytes) -> str:
    return hashlib.sha256(content).hexdigest()


@st.cache_data(show_spinner=False, max_entries=60)
def cached_load(
    name: str,
    content_hash: str,
    content: bytes,
    enabled: bool,
    min_chars: int,
    max_pages: int,
    dpi: int,
) -> list[dict]:
    del content_hash

    config = OCRConfig(
        enabled=enabled,
        min_useful_characters=min_chars,
        max_ocr_pages=max_pages,
        dpi=dpi,
    )

    return [
        trace.to_dict()
        for trace in load_document(
            name,
            content,
            config,
        )
    ]


def restore(data: dict) -> PageTrace:
    return PageTrace(**data)


def build_dossier_word(
    metadata: pd.DataFrame,
    timeline: pd.DataFrame,
    history: list[dict],
) -> bytes:
    document = Document()
    document.styles["Normal"].font.name = "Arial"
    document.styles["Normal"].font.size = Pt(10)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DOSSIER DE ESTRATEGIA JURÍDICA")
    run.bold = True
    run.font.size = Pt(15)

    document.add_heading("1. Datos del proceso", level=1)

    for _, row in metadata.iterrows():
        paragraph = document.add_paragraph()
        label = paragraph.add_run(f"{row['Campo']}: ")
        label.bold = True
        paragraph.add_run(str(row["Valor"] or ""))

    document.add_heading("2. Línea de tiempo", level=1)

    for _, row in timeline.iterrows():
        document.add_paragraph(
            f"{row.get('Fecha principal', '')} — "
            f"{row.get('Tipo', '')} — "
            f"{row.get('Documento', '')}"
        )

    document.add_heading("3. Preguntas y respuestas", level=1)

    if not history:
        document.add_paragraph(
            "No se registraron preguntas durante la sesión."
        )
    else:
        for index, item in enumerate(history, start=1):
            document.add_heading(
                f"3.{index}. {item['question']}",
                level=2,
            )
            document.add_paragraph(
                item["answer"]
            )
            document.add_paragraph(
                f"Confianza: {item['confidence']}%"
            )

            for source in item["sources"]:
                document.add_paragraph(
                    (
                        f"Fuente: {source['document']}, "
                        f"página {source['page']}. "
                        f"Coincidencia: {source['score'] * 100:.1f}%."
                    ),
                    style="List Bullet",
                )

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


with st.sidebar:
    st.header("Lectura")

    ocr_enabled = st.checkbox(
        "Aplicar OCR",
        value=True,
    )

    min_chars = st.slider(
        "Mínimo de caracteres útiles",
        20,
        300,
        80,
        10,
    )

    max_pages = st.slider(
        "Máximo de páginas OCR",
        5,
        100,
        40,
        5,
    )

    dpi = st.select_slider(
        "Resolución OCR",
        [150, 200, 220, 250, 300],
        value=220,
    )

    top_k = st.slider(
        "Fuentes por respuesta",
        1,
        8,
        5,
    )


files = st.file_uploader(
    "Sube el expediente",
    type=[
        "pdf",
        "docx",
        "txt",
        "jpg",
        "jpeg",
        "png",
        "eml",
    ],
    accept_multiple_files=True,
)

if not files:
    st.info(
        "Carga las piezas del proceso y luego formula preguntas al expediente."
    )
    st.stop()


documents: dict[str, list[PageTrace]] = {}
progress = st.progress(0)
message = st.empty()

for index, uploaded in enumerate(files):
    message.info(
        f"Preparando {uploaded.name}"
    )

    content = uploaded.getvalue()

    try:
        raw = cached_load(
            uploaded.name,
            sha256(content),
            content,
            ocr_enabled,
            min_chars,
            max_pages,
            dpi,
        )

        documents[uploaded.name] = [
            restore(item)
            for item in raw
        ]

    except Exception as error:
        st.error(
            f"No se pudo leer {uploaded.name}: {error}"
        )
        documents[uploaded.name] = []

    progress.progress(
        (index + 1) / len(files)
    )

message.empty()


fragment_index = build_fragment_index(
    documents
)

if not fragment_index:
    st.error(
        "No fue posible construir el índice de consulta."
    )
    st.stop()


metadata = pd.DataFrame(
    [
        {"Campo": key, "Valor": value}
        for key, value in extract_case_metadata(
            documents
        ).items()
    ]
)

timeline = pd.DataFrame(
    build_timeline(documents)
)

if not timeline.empty:
    timeline["_orden"] = pd.to_datetime(
        timeline["Fecha principal"],
        errors="coerce",
    )
    timeline = timeline.sort_values(
        ["_orden", "Documento"],
        na_position="last",
    ).drop(columns="_orden")


if "strategy_history" not in st.session_state:
    st.session_state.strategy_history = []


st.subheader("1. Vista estratégica")

col1, col2, col3, col4 = st.columns(4)

col1.metric(
    "Documentos",
    len(documents),
)
col2.metric(
    "Páginas",
    sum(
        len(pages)
        for pages in documents.values()
    ),
)
col3.metric(
    "Fragmentos indexados",
    len(fragment_index),
)
col4.metric(
    "Consultas realizadas",
    len(st.session_state.strategy_history),
)


st.subheader("2. Pregunta al expediente")

suggestions = suggested_questions(
    documents
)

selected_suggestion = st.selectbox(
    "Preguntas sugeridas",
    ["Escribir otra pregunta"] + suggestions,
)

default_question = (
    ""
    if selected_suggestion == "Escribir otra pregunta"
    else selected_suggestion
)

question = st.text_area(
    "Escribe tu pregunta",
    value=default_question,
    height=100,
    placeholder=(
        "Ejemplo: ¿Qué prueba existe de que los audífonos fueron entregados?"
    ),
)

ask = st.button(
    "🔍 Consultar expediente",
    type="primary",
    use_container_width=True,
)

if ask:
    if not question.strip():
        st.warning(
            "Escribe una pregunta."
        )
    else:
        result = answer_question(
            question,
            fragment_index,
            top_k=top_k,
        )

        history_item = {
            "question": question,
            "answer": result["answer"],
            "confidence": result["confidence"],
            "status": result["status"],
            "sources": result["sources"],
        }

        st.session_state.strategy_history.append(
            history_item
        )


if st.session_state.strategy_history:
    latest = st.session_state.strategy_history[-1]

    if latest["confidence"] >= 75:
        st.success(
            f"{latest['status']} — confianza {latest['confidence']}%"
        )
    elif latest["confidence"] >= 50:
        st.warning(
            f"{latest['status']} — confianza {latest['confidence']}%"
        )
    else:
        st.error(
            f"{latest['status']} — confianza {latest['confidence']}%"
        )

    st.markdown("### Respuesta documental")
    st.write(
        latest["answer"]
    )

    st.markdown("### Fuentes")

    for index, source in enumerate(
        latest["sources"],
        start=1,
    ):
        with st.expander(
            (
                f"{index}. {source['document']} — "
                f"página {source['page']} — "
                f"{source['score'] * 100:.1f}%"
            )
        ):
            st.write(
                source["fragment"]
            )
            st.caption(
                (
                    f"Método de lectura: {source['method']} | "
                    f"Referencia interna: "
                    f"{source['document']}#p{source['page']}"
                )
            )


tab1, tab2, tab3 = st.tabs(
    [
        "🗂️ Datos del proceso",
        "🕒 Cronología",
        "📚 Historial de consultas",
    ]
)


with tab1:
    st.data_editor(
        metadata,
        use_container_width=True,
        hide_index=True,
        key="strategy_metadata",
    )


with tab2:
    st.data_editor(
        timeline,
        use_container_width=True,
        hide_index=True,
        num_rows="dynamic",
        key="strategy_timeline",
    )


with tab3:
    if not st.session_state.strategy_history:
        st.info(
            "Todavía no has realizado consultas."
        )
    else:
        for index, item in enumerate(
            reversed(
                st.session_state.strategy_history
            ),
            start=1,
        ):
            with st.expander(
                f"{index}. {item['question']}"
            ):
                st.write(
                    item["answer"]
                )
                st.caption(
                    (
                        f"Confianza: {item['confidence']}% | "
                        f"{item['status']}"
                    )
                )


st.subheader("3. Dossier de estrategia")

download_col1, download_col2 = st.columns(2)

history_table = pd.DataFrame(
    [
        {
            "Pregunta": item["question"],
            "Respuesta": item["answer"],
            "Confianza": item["confidence"],
            "Estado": item["status"],
            "Fuentes": " | ".join(
                (
                    f"{source['document']}, "
                    f"página {source['page']}"
                )
                for source in item["sources"]
            ),
        }
        for item in st.session_state.strategy_history
    ]
)

excel = io.BytesIO()

with pd.ExcelWriter(
    excel,
    engine="openpyxl",
) as writer:
    metadata.to_excel(
        writer,
        sheet_name="Datos proceso",
        index=False,
    )
    timeline.to_excel(
        writer,
        sheet_name="Cronologia",
        index=False,
    )
    history_table.to_excel(
        writer,
        sheet_name="Consultas",
        index=False,
    )


with download_col1:
    st.download_button(
        "Descargar dossier en Excel",
        data=excel.getvalue(),
        file_name="dossier_estrategia_juridica.xlsx",
        mime=(
            "application/vnd.openxmlformats-officedocument."
            "spreadsheetml.sheet"
        ),
        use_container_width=True,
    )


with download_col2:
    st.download_button(
        "Descargar dossier en Word",
        data=build_dossier_word(
            metadata,
            timeline,
            st.session_state.strategy_history,
        ),
        file_name="dossier_estrategia_juridica.docx",
        mime=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
        use_container_width=True,
    )


if st.button(
    "🧹 Limpiar historial de consultas"
):
    st.session_state.strategy_history = []
    st.rerun()


st.warning(
    "La Sala de Estrategia responde únicamente con fragmentos del expediente. "
    "No reemplaza la valoración jurídica profesional y puede omitir contexto "
    "si los documentos están incompletos o tienen baja calidad."
)
