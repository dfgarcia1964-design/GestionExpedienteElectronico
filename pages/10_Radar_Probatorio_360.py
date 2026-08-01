from __future__ import annotations

import hashlib
import io

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from legal_analyzer.case_extractor import build_timeline, extract_case_metadata
from legal_analyzer.document_loader import load_document
from legal_analyzer.evidence_radar import (
    evidence_inventory,
    extract_entities,
    extract_relations,
    graphviz_source,
    missing_evidence_actions,
    theory_of_case,
)
from legal_analyzer.models import PageTrace
from legal_analyzer.ocr_engine import OCRConfig


st.set_page_config(
    page_title="Radar Probatorio 360",
    page_icon="🛰️",
    layout="wide",
)

st.title("🛰️ Radar Probatorio 360")
st.caption(
    "Mapa visual de personas, entidades, actuaciones, pruebas localizadas "
    "y vacíos probatorios del expediente."
)


def digest(content: bytes) -> str:
    return hashlib.sha256(content).hexdigest()


@st.cache_data(show_spinner=False, max_entries=60)
def cached_load(name, content_hash, content, enabled, min_chars, max_pages, dpi):
    del content_hash
    config = OCRConfig(
        enabled=enabled,
        min_useful_characters=min_chars,
        max_ocr_pages=max_pages,
        dpi=dpi,
    )
    return [item.to_dict() for item in load_document(name, content, config)]


def restore(data: dict) -> PageTrace:
    return PageTrace(**data)


def word_report(metadata, theory, inventory, relations, actions) -> bytes:
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RADAR PROBATORIO 360")
    run.bold = True
    run.font.size = Pt(15)

    doc.add_heading("1. Síntesis estratégica", level=1)
    doc.add_paragraph(theory)

    doc.add_heading("2. Datos del proceso", level=1)
    for key, value in metadata.items():
        p = doc.add_paragraph()
        r = p.add_run(f"{key}: ")
        r.bold = True
        p.add_run(str(value or ""))

    doc.add_heading("3. Inventario probatorio", level=1)
    for row in inventory:
        doc.add_paragraph(
            f"{row['Tipo de prueba']}: {row['Estado']} — {row['Fuentes']}",
            style="List Bullet",
        )

    doc.add_heading("4. Vacíos y acciones sugeridas", level=1)
    for action in actions:
        doc.add_paragraph(action, style="List Number")

    doc.add_heading("5. Relaciones documentales", level=1)
    for row in relations[:50]:
        doc.add_paragraph(
            f"{row['Origen']} {row['Relación']} {row['Destino']} "
            f"({row['Documento']}, página {row['Página']})",
            style="List Bullet",
        )

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


with st.sidebar:
    st.header("Lectura")
    enabled = st.checkbox("Aplicar OCR", value=True)
    min_chars = st.slider("Mínimo de caracteres útiles", 20, 300, 80, 10)
    max_pages = st.slider("Máximo de páginas OCR", 5, 100, 40, 5)
    dpi = st.select_slider("Resolución OCR", [150, 200, 220, 250, 300], value=220)

files = st.file_uploader(
    "Sube las piezas del expediente",
    type=["pdf", "docx", "txt", "jpg", "jpeg", "png", "eml"],
    accept_multiple_files=True,
)

if not files:
    st.info("Carga el expediente para construir el mapa probatorio.")
    st.stop()

documents = {}
bar = st.progress(0)
status = st.empty()

for index, uploaded in enumerate(files):
    status.info(f"Rastreando {uploaded.name}")
    content = uploaded.getvalue()
    try:
        raw = cached_load(
            uploaded.name,
            digest(content),
            content,
            enabled,
            min_chars,
            max_pages,
            dpi,
        )
        documents[uploaded.name] = [restore(item) for item in raw]
    except Exception as error:
        st.error(f"No se pudo procesar {uploaded.name}: {error}")
        documents[uploaded.name] = []
    bar.progress((index + 1) / len(files))

status.empty()

metadata = extract_case_metadata(documents)
entities = extract_entities(documents)
relations = extract_relations(documents, entities)
inventory = evidence_inventory(documents)
actions = missing_evidence_actions(inventory)
theory = theory_of_case(metadata, inventory, relations)

located_count = sum(row["Estado"] == "Localizada" for row in inventory)
missing_count = sum(row["Estado"] == "No localizada" for row in inventory)

c1, c2, c3, c4 = st.columns(4)
c1.metric("Entidades detectadas", len(entities))
c2.metric("Relaciones", len(relations))
c3.metric("Pruebas localizadas", located_count)
c4.metric("Vacíos probatorios", missing_count)

st.subheader("1. Mapa vivo del expediente")

if relations:
    st.graphviz_chart(
        graphviz_source(relations),
        use_container_width=True,
    )
else:
    st.warning(
        "No se detectaron relaciones suficientes para construir el mapa."
    )

tab1, tab2, tab3, tab4 = st.tabs(
    [
        "🧠 Teoría preliminar",
        "🔎 Inventario de pruebas",
        "🕳️ Vacíos probatorios",
        "🔗 Relaciones y fuentes",
    ]
)

with tab1:
    st.info(theory)
    metadata_df = pd.DataFrame(
        [{"Campo": key, "Valor": value} for key, value in metadata.items()]
    )
    st.data_editor(
        metadata_df,
        use_container_width=True,
        hide_index=True,
        key="radar_metadata",
    )

with tab2:
    inventory_df = pd.DataFrame(inventory)
    inventory_edit = st.data_editor(
        inventory_df,
        use_container_width=True,
        hide_index=True,
        column_config={
            "Estado": st.column_config.SelectboxColumn(
                "Estado",
                options=["Localizada", "No localizada", "No aplica"],
            ),
            "Revisión humana": st.column_config.TextColumn(
                "Revisión humana",
                width="large",
            ),
        },
        key="radar_inventory",
    )

with tab3:
    if actions:
        for number, action in enumerate(actions, start=1):
            st.warning(f"{number}. {action}")
    else:
        st.success(
            "No se detectaron vacíos en las categorías probatorias configuradas."
        )

with tab4:
    relations_df = pd.DataFrame(relations)
    entities_df = pd.DataFrame(entities)

    st.markdown("### Relaciones")
    st.dataframe(
        relations_df,
        use_container_width=True,
        hide_index=True,
    )

    st.markdown("### Personas y entidades")
    st.dataframe(
        entities_df,
        use_container_width=True,
        hide_index=True,
    )

st.subheader("2. Cronología probatoria")

timeline = pd.DataFrame(build_timeline(documents))
if not timeline.empty:
    timeline["_orden"] = pd.to_datetime(
        timeline["Fecha principal"],
        errors="coerce",
    )
    timeline = timeline.sort_values(
        ["_orden", "Documento"],
        na_position="last",
    ).drop(columns="_orden")

st.data_editor(
    timeline,
    use_container_width=True,
    hide_index=True,
    num_rows="dynamic",
    key="radar_timeline",
)

st.subheader("3. Exportación")

excel = io.BytesIO()
with pd.ExcelWriter(excel, engine="openpyxl") as writer:
    pd.DataFrame(
        [{"Campo": key, "Valor": value} for key, value in metadata.items()]
    ).to_excel(writer, sheet_name="Datos proceso", index=False)
    pd.DataFrame(inventory).to_excel(writer, sheet_name="Inventario pruebas", index=False)
    pd.DataFrame(actions, columns=["Acciones"]).to_excel(
        writer, sheet_name="Vacios probatorios", index=False
    )
    pd.DataFrame(relations).to_excel(writer, sheet_name="Relaciones", index=False)
    pd.DataFrame(entities).to_excel(writer, sheet_name="Entidades", index=False)
    timeline.to_excel(writer, sheet_name="Cronologia", index=False)

d1, d2 = st.columns(2)

with d1:
    st.download_button(
        "Descargar Radar 360 en Excel",
        data=excel.getvalue(),
        file_name="radar_probatorio_360.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True,
    )

with d2:
    st.download_button(
        "Descargar informe estratégico en Word",
        data=word_report(metadata, theory, inventory, relations, actions),
        file_name="informe_radar_probatorio_360.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )

st.warning(
    "El mapa y la teoría preliminar se construyen mediante coincidencias textuales. "
    "Cada relación y cada prueba deben verificarse en el documento original."
)
