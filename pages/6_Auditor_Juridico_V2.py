from __future__ import annotations

import hashlib
import io
import re
import unicodedata
from difflib import SequenceMatcher

import pandas as pd
import pytesseract
import streamlit as st
from docx import Document
from docx.shared import Pt
from pdf2image import convert_from_bytes
from PIL import ImageEnhance, ImageFilter, ImageOps
from pdf_compat import PdfReader
from pytesseract import Output

from legal_ui.system_deps import render_ocr_dependencies_status


st.set_page_config(page_title="Auditor Jurídico V2", page_icon="⚖️", layout="wide")
st.title("⚖️ Auditor jurídico V2")
st.caption(
    "Motor de cumplimiento por criterios, trazabilidad por documento y página "
    "y control de calidad del OCR."
)

render_ocr_dependencies_status(stop_if_missing=True)

STOPWORDS = {
    "para","como","esta","este","estos","estas","desde","hasta","sobre","entre",
    "dentro","ante","bajo","contra","segun","mediante","porque","cuando","donde",
    "quien","cual","del","las","los","una","unos","unas","por","con","sin","que",
    "sus","son","sea","ser","fue","han","hay","mas","al","se","de","la","el","en",
    "y","o","a","un","su"
}

VERBOS_ORDEN = (
    "ordenar","ordena","ordenese","requerir","requiere","requierase","disponer",
    "dispone","autorizar","autorice","entregar","garantizar","realizar","programar",
    "suministrar","responder","resolver","remitir","abstenerse","vincular","notificar"
)

CUMPLIMIENTO = (
    "se dio cumplimiento","dimos cumplimiento","cumplimiento del fallo","se cumplio",
    "fue entregado","se realizo","se autorizo","se programo","se remitio","se respondio",
    "se adjunta","se aporta","acta de entrega","constancia de entrega","recibido a satisfaccion"
)

INCUMPLIMIENTO = (
    "no se ha cumplido","incumplimiento","no fue entregado","no se realizo",
    "no se autorizo","no ha sido posible","pendiente","sin respuesta",
    "no existe prueba","cumplimiento parcial"
)

OPORTUNO = ("dentro del termino","dentro del plazo","oportunamente","en tiempo")
TARDIO = ("fuera del termino","extemporaneo","vencido el plazo")
INTEGRAL = ("cumplimiento integral","en su totalidad","totalmente","de manera completa")
PARCIAL = ("parcialmente","cumplimiento parcial","solo se","unicamente","queda pendiente")


def normalizar(texto: str) -> str:
    texto = texto or ""
    texto = unicodedata.normalize("NFD", texto)
    texto = "".join(c for c in texto if unicodedata.category(c) != "Mn")
    return re.sub(r"\s+", " ", texto.lower()).strip()


def limpiar(texto: str) -> str:
    texto = (texto or "").replace("\x00", " ")
    texto = re.sub(r"[ \t]+", " ", texto)
    return re.sub(r"\n{3,}", "\n\n", texto).strip()


def caracteres_utiles(texto: str) -> int:
    return len(re.sub(r"[^a-zA-ZáéíóúÁÉÍÓÚñÑ0-9]", "", texto or ""))


def palabras(texto: str) -> set[str]:
    return {
        p for p in re.findall(r"\b[a-záéíóúñ]{4,}\b", normalizar(texto))
        if p not in STOPWORDS
    }


def similitud(a: str, b: str) -> float:
    pa, pb = palabras(a), palabras(b)
    if not pa or not pb:
        return 0.0
    jaccard = len(pa & pb) / max(len(pa | pb), 1)
    secuencia = SequenceMatcher(None, " ".join(sorted(pa)), " ".join(sorted(pb))).ratio()
    return round(jaccard * 0.75 + secuencia * 0.25, 4)


def fragmentos(texto: str) -> list[str]:
    partes = re.split(
        r"(?:\n\s*\n)|(?<=[.;:])\s+(?=(?:primero|segundo|tercero|cuarto|"
        r"quinto|sexto|septimo|octavo|noveno|decimo|ordenar|ordenese|"
        r"requerir|requierase|disponer|autorizar|entregar|garantizar)\b)",
        limpiar(texto),
        flags=re.IGNORECASE,
    )
    return [re.sub(r"\s+", " ", p).strip() for p in partes if len(p.strip()) >= 30]


def etiqueta_calidad(metodo: str, confianza: float | None, utiles: int) -> str:
    if metodo != "OCR":
        if utiles >= 250:
            return "Buena"
        if utiles >= 80:
            return "Aceptable"
        return "Baja"
    confianza = confianza or 0
    if confianza >= 85 and utiles >= 150:
        return "Buena"
    if confianza >= 65 and utiles >= 80:
        return "Aceptable"
    return "Baja"


def puntaje_calidad(etiqueta: str) -> float:
    return {"Buena": 1.0, "Aceptable": 0.7, "Baja": 0.35}.get(etiqueta, 0.35)


def preprocesar(imagen):
    imagen = ImageOps.grayscale(imagen)
    imagen = ImageOps.autocontrast(imagen)
    imagen = imagen.filter(ImageFilter.MedianFilter(size=3))
    return ImageEnhance.Contrast(imagen).enhance(1.6)


def ocr_imagen(imagen) -> tuple[str, float]:
    datos = pytesseract.image_to_data(
        imagen, lang="spa", config="--oem 3 --psm 6", output_type=Output.DICT
    )
    textos, confianzas = [], []
    for texto, conf in zip(datos.get("text", []), datos.get("conf", [])):
        texto = (texto or "").strip()
        try:
            valor = float(conf)
        except (TypeError, ValueError):
            valor = -1
        if texto:
            textos.append(texto)
        if valor >= 0:
            confianzas.append(valor)
    promedio = round(sum(confianzas) / len(confianzas), 2) if confianzas else 0.0
    return limpiar(" ".join(textos)), promedio


@st.cache_data(show_spinner=False, max_entries=30)
def extraer_documento(
    nombre: str, hash_archivo: str, contenido: bytes, usar_ocr: bool,
    minimo: int, max_paginas: int, dpi: int
) -> list[dict]:
    del hash_archivo
    lector = PdfReader(io.BytesIO(contenido))
    resultado = []

    for numero, pagina in enumerate(lector.pages, start=1):
        advertencias = []
        try:
            texto = limpiar(pagina.extract_text() or "")
        except Exception as error:
            texto = ""
            advertencias.append(f"Extracción digital falló: {error}")

        utiles = caracteres_utiles(texto)
        metodo, confianza = "texto digital", None

        if usar_ocr and utiles < minimo and numero <= max_paginas:
            try:
                imagenes = convert_from_bytes(
                    contenido, dpi=dpi, first_page=numero, last_page=numero,
                    grayscale=True, thread_count=1, fmt="png"
                )
                imagen = preprocesar(imagenes[0])
                texto_ocr, confianza_ocr = ocr_imagen(imagen)

                if caracteres_utiles(texto_ocr) > utiles:
                    texto = texto_ocr
                    utiles = caracteres_utiles(texto_ocr)
                    metodo = "OCR"
                    confianza = confianza_ocr
                else:
                    advertencias.append("El OCR no mejoró el texto digital.")
            except Exception as error:
                advertencias.append(f"OCR falló: {error}")

        elif usar_ocr and utiles < minimo and numero > max_paginas:
            advertencias.append("No se aplicó OCR por el límite configurado.")

        calidad = etiqueta_calidad(metodo, confianza, utiles)
        resultado.append({
            "documento": nombre,
            "pagina": numero,
            "texto": texto,
            "metodo": metodo,
            "confianza": confianza,
            "utiles": utiles,
            "calidad": calidad,
            "advertencias": advertencias,
        })

    return resultado


def parece_fallo(paginas: list[dict]) -> bool:
    texto = normalizar(" ".join(p["texto"] for p in paginas))
    señales = (
        "fallo de tutela","administrando justicia","en merito de lo expuesto",
        "resuelve","amparar","negar el amparo"
    )
    return sum(s in texto for s in señales) >= 2


def responsable(orden: str) -> str:
    for patron in (
        r"ordenar\s+a\s+(.{3,160}?)(?:\s+que\s+|\s+para\s+|,|\.)",
        r"ordenese\s+a\s+(.{3,160}?)(?:\s+que\s+|\s+para\s+|,|\.)",
        r"requerir\s+a\s+(.{3,160}?)(?:\s+que\s+|\s+para\s+|,|\.)",
        r"requierase\s+a\s+(.{3,160}?)(?:\s+que\s+|\s+para\s+|,|\.)",
    ):
        coincidencia = re.search(patron, normalizar(orden))
        if coincidencia:
            return coincidencia.group(1).strip(" ,.;:")
    return "Requiere identificación manual"


def plazo(orden: str) -> str:
    for patron in (
        r"(?:dentro de|en el termino de|plazo de)\s+(?:las\s+|los\s+)?"
        r"(?:\d+|[a-záéíóúñ]+)\s+(?:horas|dias)(?:\s+habiles)?",
        r"termino improrrogable de\s+(?:\d+|[a-záéíóúñ]+)\s+"
        r"(?:horas|dias)(?:\s+habiles)?",
        r"de manera inmediata",
        r"inmediatamente",
    ):
        coincidencia = re.search(patron, normalizar(orden))
        if coincidencia:
            return coincidencia.group(0)
    return "No detectado"


def extraer_ordenes(nombre: str, paginas: list[dict]) -> list[dict]:
    ordenes, en_resuelve, numero = [], False, 1
    for pagina in paginas:
        texto_norm = normalizar(pagina["texto"])
        if "resuelve" in texto_norm:
            en_resuelve = True
        if not en_resuelve:
            continue
        for frag in fragmentos(pagina["texto"]):
            if any(v in normalizar(frag) for v in VERBOS_ORDEN):
                ordenes.append({
                    "id": numero,
                    "texto": frag[:2200],
                    "documento": nombre,
                    "pagina": pagina["pagina"],
                    "responsable": responsable(frag),
                    "plazo": plazo(frag),
                })
                numero += 1
    return ordenes[:40]


def señal(texto: str, positivos: tuple[str, ...], negativos: tuple[str, ...]) -> float:
    contenido = normalizar(texto)
    pos = sum(p in contenido for p in positivos)
    neg = sum(n in contenido for n in negativos)
    if pos and not neg:
        return 1.0
    if neg and not pos:
        return 0.0
    if pos and neg:
        return 0.5
    return 0.35


def señal_responsable(resp: str, frag: str) -> float:
    if "requiere identificacion manual" in normalizar(resp):
        return 0.5
    pr, pf = palabras(resp), palabras(frag)
    if not pr:
        return 0.5
    return min(1.0, len(pr & pf) / max(len(pr), 1))


def fuerza_prueba(frag: str) -> float:
    contenido = normalizar(frag)
    señales = (
        "anexo","adjunto","acta","constancia","certificacion","historia clinica",
        "factura","recibo","correo","radicado","captura","comprobante","concepto","informe"
    )
    hallazgos = sum(s in contenido for s in señales)
    return min(1.0, hallazgos * 0.12 + señal(frag, CUMPLIMIENTO, INCUMPLIMIENTO) * 0.64)


def evaluar(orden: dict, paginas_evidencia: list[dict]) -> dict:
    candidatos = []
    for pagina in paginas_evidencia:
        for indice, frag in enumerate(fragmentos(pagina["texto"]), start=1):
            sim = similitud(orden["texto"], frag)
            if sim < 0.045:
                continue
            calidad = puntaje_calidad(pagina["calidad"])
            candidatos.append({
                "fragmento": frag[:1200],
                "documento": pagina["documento"],
                "pagina": pagina["pagina"],
                "traza": f"{pagina['documento']}#p{pagina['pagina']}#f{indice}",
                "similitud": sim,
                "conducta": min(1.0, sim / 0.22),
                "responsable": señal_responsable(orden["responsable"], frag),
                "prueba": fuerza_prueba(frag),
                "oportunidad": señal(frag, OPORTUNO, TARDIO),
                "integralidad": señal(frag, INTEGRAL, PARCIAL),
                "calidad": calidad,
            })

    if not candidatos:
        return {
            "estado": "No verificable", "total": 0.0, "razon": "No se localizó prueba relacionada.",
            "evidencia": "", "traza": "Sin prueba localizada",
            "conducta": 0.0, "responsable": 0.0, "prueba": 0.0,
            "oportunidad": 0.0, "integralidad": 0.0, "calidad": 0.0,
        }

    candidatos.sort(
        key=lambda c: (
            c["conducta"] * 0.30 + c["responsable"] * 0.15 +
            c["prueba"] * 0.25 + c["oportunidad"] * 0.10 +
            c["integralidad"] * 0.10 + c["calidad"] * 0.10
        ),
        reverse=True,
    )
    mejor = candidatos[0]
    total = round(
        mejor["conducta"] * 0.30 + mejor["responsable"] * 0.15 +
        mejor["prueba"] * 0.25 + mejor["oportunidad"] * 0.10 +
        mejor["integralidad"] * 0.10 + mejor["calidad"] * 0.10,
        4,
    )

    if mejor["prueba"] < 0.28:
        estado = "No verificable"
    elif total >= 0.78 and mejor["integralidad"] >= 0.65 and mejor["responsable"] >= 0.55:
        estado = "Posible cumplimiento integral"
    elif total >= 0.60:
        estado = (
            "Posible cumplimiento parcial"
            if mejor["integralidad"] < 0.55 or mejor["oportunidad"] < 0.45
            else "Posible cumplimiento"
        )
    elif total >= 0.42:
        estado = "Requiere revisión"
    else:
        estado = "Posible incumplimiento"

    razon = (
        f"Conducta {mejor['conducta']*100:.0f}% | Responsable {mejor['responsable']*100:.0f}% | "
        f"Prueba {mejor['prueba']*100:.0f}% | Oportunidad {mejor['oportunidad']*100:.0f}% | "
        f"Integralidad {mejor['integralidad']*100:.0f}% | Calidad {mejor['calidad']*100:.0f}%"
    )
    return {
        "estado": estado, "total": total, "razon": razon,
        "evidencia": mejor["fragmento"],
        "traza": f"{mejor['documento']}, página {mejor['pagina']}, {mejor['traza']}",
        **{k: mejor[k] for k in ("conducta","responsable","prueba","oportunidad","integralidad","calidad")},
    }


def informe_word(matriz: pd.DataFrame, calidad: pd.DataFrame) -> bytes:
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)
    doc.add_heading("AUDITORÍA DE CUMPLIMIENTO DEL FALLO", 0)
    doc.add_paragraph(
        "Resultado preliminar. Debe contrastarse con el expediente original y revisarse jurídicamente."
    )
    for _, fila in matriz.iterrows():
        doc.add_heading(f"Orden {fila['N.º']} — {fila['Estado revisado']}", level=2)
        for campo in (
            "Orden judicial","Fuente de la orden","Responsable","Plazo","Prueba localizada",
            "Trazabilidad","Puntaje total","Razón automática","Observaciones humanas"
        ):
            p = doc.add_paragraph()
            r = p.add_run(f"{campo}: ")
            r.bold = True
            p.add_run(str(fila.get(campo, "")))
    doc.add_heading("Control de calidad OCR", level=1)
    for _, fila in calidad.iterrows():
        doc.add_paragraph(
            f"{fila['Documento']}, página {fila['Página']}: {fila['Método']} — "
            f"calidad {fila['Calidad']} — confianza {fila['Confianza OCR']}"
        )
    salida = io.BytesIO()
    doc.save(salida)
    return salida.getvalue()


with st.sidebar:
    st.header("Configuración OCR")
    usar_ocr = st.checkbox("Aplicar OCR cuando falte texto", value=True)
    minimo = st.slider("Mínimo de caracteres útiles", 20, 300, 80, 10)
    max_paginas = st.slider("Máximo de páginas OCR por archivo", 5, 100, 40, 5)
    dpi = st.select_slider("Resolución OCR", options=[150,200,220,250,300], value=220)

archivos = st.file_uploader(
    "Sube el fallo y las piezas posteriores", type=["pdf"], accept_multiple_files=True
)

if not archivos:
    st.info("Carga el fallo, respuestas, actas, constancias, conceptos y requerimientos.")
    st.stop()

documentos = {}
barra = st.progress(0)
estado = st.empty()

for i, archivo in enumerate(archivos):
    estado.info(f"Procesando {archivo.name}")
    contenido = archivo.getvalue()
    digest = hashlib.sha256(contenido).hexdigest()
    try:
        documentos[archivo.name] = extraer_documento(
            archivo.name, digest, contenido, usar_ocr, minimo, max_paginas, dpi
        )
    except Exception as error:
        st.error(f"No se pudo procesar {archivo.name}: {error}")
        documentos[archivo.name] = []
    barra.progress((i + 1) / len(archivos))
estado.empty()

filas_calidad = []
for nombre, paginas in documentos.items():
    for p in paginas:
        filas_calidad.append({
            "Documento": nombre,
            "Página": p["pagina"],
            "Método": p["metodo"],
            "Confianza OCR": "" if p["confianza"] is None else f"{p['confianza']:.1f}%",
            "Caracteres útiles": p["utiles"],
            "Calidad": p["calidad"],
            "Advertencias": " | ".join(p["advertencias"]),
            "Vista previa": p["texto"][:300],
        })

tabla_calidad = pd.DataFrame(filas_calidad)
st.subheader("1. Control de calidad del OCR")
st.dataframe(tabla_calidad, use_container_width=True, hide_index=True)

c1, c2, c3 = st.columns(3)
c1.metric("Páginas", len(tabla_calidad))
c2.metric("Páginas OCR", int((tabla_calidad["Método"] == "OCR").sum()))
c3.metric("Calidad baja", int((tabla_calidad["Calidad"] == "Baja").sum()))

if (tabla_calidad["Calidad"] == "Baja").any():
    st.warning("Hay páginas de calidad baja. Verifica las conclusiones contra la imagen original.")

nombres = list(documentos)
candidatos = [n for n, p in documentos.items() if parece_fallo(p)]
indice = nombres.index(candidatos[0]) if candidatos else 0
fallo = st.selectbox("Selecciona el documento que contiene el fallo", nombres, index=indice)

ordenes = extraer_ordenes(fallo, documentos[fallo])
if not ordenes:
    st.error("No se detectaron órdenes. Revisa el fallo seleccionado y la calidad del OCR.")
    st.stop()

st.subheader("2. Órdenes y trazabilidad")
tabla_ordenes = pd.DataFrame([{
    "N.º": o["id"], "Orden judicial": o["texto"],
    "Fuente": f"{o['documento']}, página {o['pagina']}",
    "Responsable": o["responsable"], "Plazo": o["plazo"]
} for o in ordenes])
st.dataframe(tabla_ordenes, use_container_width=True, hide_index=True)

paginas_evidencia = [
    p for nombre, paginas in documentos.items() if nombre != fallo for p in paginas
]
filas = []
for o in ordenes:
    a = evaluar(o, paginas_evidencia)
    filas.append({
        "N.º": o["id"],
        "Orden judicial": o["texto"],
        "Fuente de la orden": f"{o['documento']}, página {o['pagina']}",
        "Responsable": o["responsable"],
        "Plazo": o["plazo"],
        "Prueba localizada": a["evidencia"],
        "Trazabilidad": a["traza"],
        "Coincidencia de conducta": f"{a['conducta']*100:.0f}%",
        "Responsable correcto": f"{a['responsable']*100:.0f}%",
        "Fuerza de la prueba": f"{a['prueba']*100:.0f}%",
        "Oportunidad": f"{a['oportunidad']*100:.0f}%",
        "Integralidad": f"{a['integralidad']*100:.0f}%",
        "Calidad OCR/texto": f"{a['calidad']*100:.0f}%",
        "Puntaje total": f"{a['total']*100:.1f}%",
        "Estado automático": a["estado"],
        "Estado revisado": a["estado"],
        "Razón automática": a["razon"],
        "Observaciones humanas": "",
    })

matriz = pd.DataFrame(filas)
estados = [
    "Cumplida","Parcialmente cumplida","Incumplida","No verificable",
    "Requiere revisión","Posible cumplimiento integral","Posible cumplimiento",
    "Posible cumplimiento parcial","Posible incumplimiento"
]

st.subheader("3. Matriz de cumplimiento por criterios")
editada = st.data_editor(
    matriz, use_container_width=True, hide_index=True, num_rows="dynamic",
    column_config={
        "Estado revisado": st.column_config.SelectboxColumn(
            "Estado revisado", options=estados, required=True
        ),
        "Orden judicial": st.column_config.TextColumn("Orden judicial", width="large"),
        "Prueba localizada": st.column_config.TextColumn("Prueba localizada", width="large"),
        "Razón automática": st.column_config.TextColumn("Razón automática", width="large"),
        "Observaciones humanas": st.column_config.TextColumn("Observaciones humanas", width="large"),
    },
    disabled=[
        "N.º","Fuente de la orden","Coincidencia de conducta","Responsable correcto",
        "Fuerza de la prueba","Oportunidad","Integralidad","Calidad OCR/texto",
        "Puntaje total","Estado automático","Razón automática"
    ],
    key="auditor_v2"
)

st.subheader("4. Resumen y exportación")
resumen = editada["Estado revisado"].fillna("Sin estado").value_counts().rename_axis("Estado").reset_index(name="Cantidad")
st.dataframe(resumen, use_container_width=True, hide_index=True)

excel = io.BytesIO()
with pd.ExcelWriter(excel, engine="openpyxl") as escritor:
    editada.to_excel(escritor, sheet_name="Cumplimiento", index=False)
    tabla_calidad.to_excel(escritor, sheet_name="Calidad OCR", index=False)
    tabla_ordenes.to_excel(escritor, sheet_name="Ordenes", index=False)

d1, d2 = st.columns(2)
with d1:
    st.download_button(
        "Descargar auditoría en Excel", data=excel.getvalue(),
        file_name="auditoria_cumplimiento_v2.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True
    )
with d2:
    st.download_button(
        "Descargar auditoría en Word", data=informe_word(editada, tabla_calidad),
        file_name="auditoria_cumplimiento_v2.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )

st.warning(
    "El motor no declara por sí solo un cumplimiento jurídico definitivo. "
    "El estado revisado debe basarse en la orden completa, la prueba original, "
    "la fecha y la identidad del responsable."
)
