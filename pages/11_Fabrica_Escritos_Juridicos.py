from __future__ import annotations

import io

import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from legal_analyzer.legal_drafter import DOCUMENT_TITLES, build_legal_text


st.set_page_config(
    page_title="Fábrica de Escritos Jurídicos",
    page_icon="📝",
    layout="wide",
)

st.title("📝 Fábrica de Escritos Jurídicos")
st.caption(
    "Convierte los hallazgos del expediente en un escrito editable y descargable en Word."
)


def create_word(data: dict[str, str]) -> bytes:
    document = Document()
    document.styles["Normal"].font.name = "Arial"
    document.styles["Normal"].font.size = Pt(11)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(data["title"])
    run.bold = True
    run.font.size = Pt(14)

    document.add_paragraph(data["heading"])
    document.add_paragraph(data["introduction"])

    document.add_heading("I. HECHOS", level=1)
    document.add_paragraph(data["facts"])

    document.add_heading("II. SOLICITUDES", level=1)
    document.add_paragraph(data["requests"])

    document.add_heading("III. PRUEBAS Y ANEXOS", level=1)
    document.add_paragraph(data["evidence"])

    document.add_heading("IV. FUNDAMENTOS", level=1)
    document.add_paragraph(data["legal_basis"])

    if data["observations"]:
        document.add_heading("V. OBSERVACIONES", level=1)
        document.add_paragraph(data["observations"])

    document.add_paragraph(data["closing"])

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


document_type = st.selectbox(
    "Tipo de escrito",
    list(DOCUMENT_TITLES.keys()),
)

col1, col2 = st.columns(2)

with col1:
    city = st.text_input("Ciudad", value="Popayán")
    court = st.text_input("Juzgado o autoridad")
    claimant = st.text_input("Accionante o solicitante")

with col2:
    respondent = st.text_input("Entidad accionada o destinataria")
    case_number = st.text_input("Radicado")
    email = st.text_input("Correo de notificación")

st.subheader("Contenido del escrito")

facts = st.text_area(
    "Hechos",
    height=220,
    placeholder=(
        "1. El día ...\n"
        "2. Mediante fallo de fecha ... el juzgado ordenó ...\n"
        "3. La entidad informó ...\n"
        "4. Sin embargo, en el documento ... página ... se observa ..."
    ),
)

requests = st.text_area(
    "Solicitudes concretas",
    height=170,
    placeholder=(
        "1. Requerir a la entidad para que...\n"
        "2. Verificar materialmente...\n"
        "3. Incorporar al expediente..."
    ),
)

evidence = st.text_area(
    "Pruebas y anexos",
    height=150,
    placeholder=(
        "Anexo 1. Fallo de tutela, página...\n"
        "Anexo 2. Respuesta de la entidad, página...\n"
        "Anexo 3. Constancia de envío..."
    ),
)

legal_basis = st.text_area(
    "Fundamentos jurídicos",
    height=150,
    placeholder=(
        "Incluye únicamente normas y jurisprudencia verificadas para el caso."
    ),
)

observations = st.text_area(
    "Observaciones adicionales",
    height=120,
)

data = build_legal_text(
    document_type=document_type,
    city=city,
    court=court,
    claimant=claimant,
    respondent=respondent,
    case_number=case_number,
    facts=facts,
    requests=requests,
    evidence=evidence,
    legal_basis=legal_basis,
    observations=observations,
)

st.subheader("Vista previa")

preview = (
    f"{data['heading']}\n\n"
    f"{data['introduction']}\n\n"
    f"I. HECHOS\n{data['facts']}\n\n"
    f"II. SOLICITUDES\n{data['requests']}\n\n"
    f"III. PRUEBAS Y ANEXOS\n{data['evidence']}\n\n"
    f"IV. FUNDAMENTOS\n{data['legal_basis']}\n\n"
    f"{data['closing']}"
)

st.text_area(
    "Documento generado",
    value=preview,
    height=600,
)

file_name = (
    document_type.lower()
    .replace(" ", "_")
    .replace("ó", "o")
    .replace("í", "i")
    .replace("é", "e")
    .replace("á", "a")
    .replace("ú", "u")
    .replace("ñ", "n")
    + ".docx"
)

st.download_button(
    "Descargar escrito en Word",
    data=create_word(data),
    file_name=file_name,
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    type="primary",
    use_container_width=True,
)

st.warning(
    "El documento es un borrador editable. Verifica hechos, fechas, anexos, "
    "competencia, términos, normas y jurisprudencia antes de presentarlo."
)
