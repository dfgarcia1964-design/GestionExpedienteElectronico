from __future__ import annotations

import hashlib
import io

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from legal_analyzer.case_extractor import (
    build_timeline,
    classify_document,
    extract_case_metadata,
)
from legal_analyzer.command_center import (
    assess_risk,
    detect_contradictions,
    executive_summary,
)
from legal_analyzer.document_loader import load_document
from legal_analyzer.models import PageTrace
from legal_analyzer.ocr_engine import OCRConfig, quality_label


st.set_page_config(
    page_title="Centro de Mando Jurídico",
    page_icon="🚨",
    layout="wide",
)

st.title("🚨 Centro de Mando Jurídico")
st.caption(
    "Panel ejecutivo del expediente: riesgo, contradicciones, documentos faltantes "
    "y acciones prioritarias."
)


def sha256(content: bytes) -> str:
    return hashlib.sha256(content).hexdigest()


@st.cache_data(show_spinner=False, max_entries=60)
def cached_load(
    name: str,
    content_hash: str,
    content: bytes,
    enabled: bool,
    min_chars: int,
    max_pages: int,
    dpi: int,
) -> list[dict]:
    del content_hash

    config = OCRConfig(
        enabled=enabled,
        min_useful_characters=min_chars,
        max_ocr_pages=max_pages,
        dpi=dpi,
    )

    return [
        trace.to_dict()
        for trace in load_document(name, content, config)
    ]


def restore(data: dict) -> PageTrace:
    return PageTrace(**data)


def build_word_report(
    metadata: pd.DataFrame,
    timeline: pd.DataFrame,
    contradictions: pd.DataFrame,
    risk_score: int,
    risk_level: str,
    summary: str,
    reasons: list[str],
    actions: list[str],
) -> bytes:
    document = Document()
    document.styles["Normal"].font.name = "Arial"
    document.styles["Normal"].font.size = Pt(10)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CENTRO DE MANDO JURÍDICO")
    run.bold = True
    run.font.size = Pt(15)

    document.add_paragraph(summary)

    document.add_heading("1. Nivel de riesgo", level=1)
    document.add_paragraph(
        f"{risk_level} — {risk_score}/100"
    )

    document.add_heading("2. Razones del riesgo", level=1)
    for reason in reasons:
        document.add_paragraph(reason, style="List Bullet")

    document.add_heading("3. Acciones prioritarias", level=1)
    for action in actions:
        document.add_paragraph(action, style="List Number")

    document.add_heading("4. Datos del proceso", level=1)
    for _, row in metadata.iterrows():
        paragraph = document.add_paragraph()
        label = paragraph.add_run(f"{row['Campo']}: ")
        label.bold = True
        paragraph.add_run(str(row["Valor"] or ""))

    document.add_heading("5. Línea de tiempo", level=1)
    for _, row in timeline.iterrows():
        document.add_paragraph(
            f"{row.get('Fecha principal', '')} — "
            f"{row.get('Tipo', '')} — "
            f"{row.get('Documento', '')}"
        )

    document.add_heading("6. Contradicciones potenciales", level=1)

    if contradictions.empty:
        document.add_paragraph(
            "No se detectaron contradicciones con las reglas actuales."
        )
    else:
        for index, row in contradictions.iterrows():
            document.add_heading(
                f"Contradicción {index + 1}: {row.get('Tema', '')}",
                level=2,
            )
            document.add_paragraph(
                f"Versión 1: {row.get('Versión 1', '')}"
            )
            document.add_paragraph(
                f"Fuente 1: {row.get('Fuente 1', '')}"
            )
            document.add_paragraph(
                f"Versión 2: {row.get('Versión 2', '')}"
            )
            document.add_paragraph(
                f"Fuente 2: {row.get('Fuente 2', '')}"
            )

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


with st.sidebar:
    st.header("Configuración")

    ocr_enabled = st.checkbox(
        "Aplicar OCR",
        value=True,
    )

    min_chars = st.slider(
        "Mínimo de caracteres útiles",
        20,
        300,
        80,
        10,
    )

    max_pages = st.slider(
        "Máximo de páginas OCR por PDF",
        5,
        100,
        40,
        5,
    )

    dpi = st.select_slider(
        "Resolución OCR",
        [150, 200, 220, 250, 300],
        value=220,
    )


files = st.file_uploader(
    "Sube el expediente completo o una selección de piezas",
    type=["pdf", "docx", "txt", "jpg", "jpeg", "png", "eml"],
    accept_multiple_files=True,
)

if not files:
    st.info(
        "Carga tutela, fallo, respuestas, autos, constancias, desacatos, "
        "derechos de petición y anexos."
    )
    st.stop()


documents: dict[str, list[PageTrace]] = {}
progress = st.progress(0)
message = st.empty()

for index, uploaded in enumerate(files):
    message.info(f"Analizando {uploaded.name}")

    content = uploaded.getvalue()

    try:
        raw = cached_load(
            uploaded.name,
            sha256(content),
            content,
            ocr_enabled,
            min_chars,
            max_pages,
            dpi,
        )

        documents[uploaded.name] = [
            restore(item)
            for item in raw
        ]

    except Exception as error:
        st.error(
            f"No se pudo procesar {uploaded.name}: {error}"
        )
        documents[uploaded.name] = []

    progress.progress((index + 1) / len(files))

message.empty()


quality_rows = []

for name, pages in documents.items():
    kind, _ = classify_document(pages)

    for page in pages:
        quality_rows.append(
            {
                "Documento": name,
                "Tipo": kind,
                "Página": page.page,
                "Método": page.extraction_method,
                "Confianza OCR": (
                    ""
                    if page.ocr_confidence is None
                    else f"{page.ocr_confidence:.1f}%"
                ),
                "Caracteres útiles": page.useful_characters,
                "Calidad": quality_label(page),
                "Advertencias": " | ".join(page.warnings),
            }
        )

quality = pd.DataFrame(quality_rows)
contradictions = pd.DataFrame(
    detect_contradictions(documents)
)
risk = assess_risk(
    documents,
    quality_rows,
)
summary = executive_summary(
    documents,
    risk,
    contradictions.to_dict("records"),
)


st.subheader("1. Estado general del expediente")

col1, col2, col3, col4 = st.columns(4)

col1.metric(
    "Riesgo",
    risk.level,
    f"{risk.score}/100",
)
col2.metric(
    "Documentos",
    len(documents),
)
col3.metric(
    "Contradicciones",
    len(contradictions),
)
col4.metric(
    "Páginas de baja calidad",
    int(
        (quality["Calidad"] == "Baja").sum()
        if not quality.empty
        else 0
    ),
)

st.progress(risk.score / 100)

if risk.level in {"Crítico", "Alto"}:
    st.error(summary)
elif risk.level == "Medio":
    st.warning(summary)
else:
    st.success(summary)


tab1, tab2, tab3, tab4, tab5 = st.tabs(
    [
        "🎯 Qué hacer ahora",
        "⚠️ Riesgos",
        "🧩 Contradicciones",
        "🕒 Línea de tiempo",
        "📄 Calidad documental",
    ]
)


with tab1:
    st.subheader("Acciones prioritarias")

    for number, action in enumerate(
        risk.next_actions,
        start=1,
    ):
        st.markdown(f"**{number}. {action}**")

    st.info(
        "Estas son opciones de revisión y actuación, no una orden jurídica automática."
    )


with tab2:
    st.subheader("Razones del nivel de riesgo")

    for reason in risk.reasons:
        st.warning(reason)


with tab3:
    if contradictions.empty:
        st.success(
            "No se detectaron contradicciones con las reglas actuales."
        )
    else:
        st.dataframe(
            contradictions,
            use_container_width=True,
            hide_index=True,
        )


with tab4:
    timeline = pd.DataFrame(
        build_timeline(documents)
    )

    if not timeline.empty:
        timeline["_orden"] = pd.to_datetime(
            timeline["Fecha principal"],
            errors="coerce",
        )
        timeline = timeline.sort_values(
            ["_orden", "Documento"],
            na_position="last",
        ).drop(columns="_orden")

    timeline_edit = st.data_editor(
        timeline,
        use_container_width=True,
        hide_index=True,
        num_rows="dynamic",
        key="command_timeline",
    )


with tab5:
    st.dataframe(
        quality,
        use_container_width=True,
        hide_index=True,
    )


st.subheader("2. Datos del proceso")

metadata = pd.DataFrame(
    [
        {"Campo": key, "Valor": value}
        for key, value in extract_case_metadata(
            documents
        ).items()
    ]
)

metadata_edit = st.data_editor(
    metadata,
    use_container_width=True,
    hide_index=True,
    column_config={
        "Campo": st.column_config.TextColumn(
            "Campo",
            disabled=True,
        ),
        "Valor": st.column_config.TextColumn(
            "Valor",
            width="large",
        ),
    },
    key="command_metadata",
)


st.subheader("3. Exportación ejecutiva")

excel = io.BytesIO()

with pd.ExcelWriter(
    excel,
    engine="openpyxl",
) as writer:
    metadata_edit.to_excel(
        writer,
        sheet_name="Datos proceso",
        index=False,
    )
    timeline_edit.to_excel(
        writer,
        sheet_name="Linea de tiempo",
        index=False,
    )
    quality.to_excel(
        writer,
        sheet_name="Calidad documental",
        index=False,
    )
    contradictions.to_excel(
        writer,
        sheet_name="Contradicciones",
        index=False,
    )
    pd.DataFrame(
        {
            "Razones": risk.reasons,
        }
    ).to_excel(
        writer,
        sheet_name="Riesgos",
        index=False,
    )
    pd.DataFrame(
        {
            "Acciones": risk.next_actions,
        }
    ).to_excel(
        writer,
        sheet_name="Acciones",
        index=False,
    )


download1, download2 = st.columns(2)

with download1:
    st.download_button(
        "Descargar centro de mando en Excel",
        data=excel.getvalue(),
        file_name="centro_mando_juridico.xlsx",
        mime=(
            "application/vnd.openxmlformats-officedocument."
            "spreadsheetml.sheet"
        ),
        use_container_width=True,
    )

with download2:
    st.download_button(
        "Descargar informe ejecutivo en Word",
        data=build_word_report(
            metadata_edit,
            timeline_edit,
            contradictions,
            risk.score,
            risk.level,
            summary,
            risk.reasons,
            risk.next_actions,
        ),
        file_name="informe_centro_mando_juridico.docx",
        mime=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
        use_container_width=True,
    )


st.warning(
    "El semáforo de riesgo y las contradicciones son señales automáticas. "
    "Debes validar cada conclusión con el documento y la página originales."
)
