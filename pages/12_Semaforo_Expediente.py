from __future__ import annotations

import hashlib
import io

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from legal_analyzer.command_center import detect_contradictions
from legal_analyzer.document_loader import load_document
from legal_analyzer.models import PageTrace
from legal_analyzer.ocr_engine import OCRConfig, quality_label
from legal_analyzer.semaphore_engine import (
    build_semaphores,
    overall_semaphore,
)


st.set_page_config(
    page_title="Semáforo del Expediente",
    page_icon="🚦",
    layout="wide",
)

st.title("🚦 Semáforo del Expediente")
st.caption(
    "Identifica qué está controlado, qué debes revisar y qué requiere atención urgente."
)


def sha256(content: bytes) -> str:
    return hashlib.sha256(content).hexdigest()


@st.cache_data(show_spinner=False, max_entries=60)
def cached_load(name, content_hash, content, enabled, min_chars, max_pages, dpi):
    del content_hash
    config = OCRConfig(
        enabled=enabled,
        min_useful_characters=min_chars,
        max_ocr_pages=max_pages,
        dpi=dpi,
    )
    return [trace.to_dict() for trace in load_document(name, content, config)]


def restore(data: dict) -> PageTrace:
    return PageTrace(**data)


def badge(color: str) -> str:
    icons = {"Verde": "🟢", "Amarillo": "🟡", "Rojo": "🔴"}
    return f"{icons.get(color, '⚪')} {color}"


def create_word(items, overall, contradictions_df) -> bytes:
    document = Document()
    document.styles["Normal"].font.name = "Arial"
    document.styles["Normal"].font.size = Pt(10)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SEMÁFORO DEL EXPEDIENTE")
    run.bold = True
    run.font.size = Pt(15)

    document.add_paragraph(
        f"Estado general: {overall['label']} — "
        f"{overall['color']} — {overall['score']}/100"
    )

    for item in items:
        document.add_heading(
            f"{item.area} — {item.color}",
            level=1,
        )
        document.add_paragraph(f"Nivel de atención: {item.score}/100")
        document.add_paragraph(f"Motivo: {item.reason}")
        document.add_paragraph(f"Acción sugerida: {item.action}")
        if item.source:
            document.add_paragraph(f"Fuentes: {item.source}")

    document.add_heading("CONTRADICCIONES COMPLETAS", level=1)

    if contradictions_df.empty:
        document.add_paragraph(
            "No se detectaron contradicciones con las reglas actuales."
        )
    else:
        for index, row in contradictions_df.iterrows():
            document.add_heading(
                f"Contradicción {index + 1}: {row.get('Tema', '')}",
                level=2,
            )
            document.add_paragraph(
                f"Primera versión completa:\n{row.get('Versión 1 completa', '')}"
            )
            document.add_paragraph(
                f"Fuente 1: {row.get('Fuente 1', '')}"
            )
            document.add_paragraph(
                f"Segunda versión completa:\n{row.get('Versión 2 completa', '')}"
            )
            document.add_paragraph(
                f"Fuente 2: {row.get('Fuente 2', '')}"
            )
            document.add_paragraph(
                f"Tipo de oposición: {row.get('Tipo de oposición', '')}"
            )
            document.add_paragraph(
                f"Conclusión revisada: {row.get('Conclusión revisada', '')}"
            )
            document.add_paragraph(
                f"Observaciones: {row.get('Observaciones', '')}"
            )

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


with st.sidebar:
    st.header("Lectura")
    ocr_enabled = st.checkbox("Aplicar OCR", value=True)
    min_chars = st.slider("Mínimo de caracteres útiles", 20, 300, 80, 10)
    max_pages = st.slider("Máximo de páginas OCR", 5, 100, 40, 5)
    dpi = st.select_slider(
        "Resolución OCR",
        [150, 200, 220, 250, 300],
        value=220,
    )

files = st.file_uploader(
    "Sube las piezas del expediente",
    type=["pdf", "docx", "txt", "jpg", "jpeg", "png", "eml"],
    accept_multiple_files=True,
)

if not files:
    st.info("Carga el expediente para generar los semáforos.")
    st.stop()

documents = {}
progress = st.progress(0)
message = st.empty()

for index, uploaded in enumerate(files):
    message.info(f"Evaluando {uploaded.name}")
    content = uploaded.getvalue()

    try:
        raw = cached_load(
            uploaded.name,
            sha256(content),
            content,
            ocr_enabled,
            min_chars,
            max_pages,
            dpi,
        )
        documents[uploaded.name] = [restore(item) for item in raw]
    except Exception as error:
        st.error(f"No se pudo procesar {uploaded.name}: {error}")
        documents[uploaded.name] = []

    progress.progress((index + 1) / len(files))

message.empty()

quality_rows = []
for name, pages in documents.items():
    for page in pages:
        quality_rows.append(
            {
                "Documento": name,
                "Página": page.page,
                "Calidad": quality_label(page),
                "Método": page.extraction_method,
                "Confianza OCR": page.ocr_confidence,
            }
        )

items = build_semaphores(documents, quality_rows)
overall = overall_semaphore(items)
contradictions = detect_contradictions(documents)
contradictions_df = pd.DataFrame(contradictions)

st.subheader("1. Estado general")
col1, col2, col3 = st.columns([1.3, 1, 1])
col1.metric("Estado", overall["label"])
col2.metric("Semáforo", badge(overall["color"]))
col3.metric("Nivel de atención", f"{overall['score']}/100")
st.progress(overall["score"] / 100)

if overall["color"] == "Rojo":
    st.error("Hay aspectos críticos que requieren atención prioritaria.")
elif overall["color"] == "Amarillo":
    st.warning("El expediente tiene alertas que deben revisarse antes de actuar.")
else:
    st.success("El expediente está razonablemente controlado con la información cargada.")

st.subheader("2. Semáforos por área")
columns = st.columns(3)

for index, item in enumerate(items):
    with columns[index % 3]:
        if item.color == "Rojo":
            st.error(
                f"### 🔴 {item.area}\n"
                f"**Atención:** {item.score}/100\n\n"
                f"**Qué cuidar:** {item.reason}\n\n"
                f"**Qué hacer:** {item.action}"
            )
        elif item.color == "Amarillo":
            st.warning(
                f"### 🟡 {item.area}\n"
                f"**Atención:** {item.score}/100\n\n"
                f"**Qué revisar:** {item.reason}\n\n"
                f"**Qué hacer:** {item.action}"
            )
        else:
            st.success(
                f"### 🟢 {item.area}\n"
                f"**Atención:** {item.score}/100\n\n"
                f"**Controlado:** {item.reason}\n\n"
                f"**Mantener:** {item.action}"
            )

st.subheader("3. Contradicciones completas")

if contradictions_df.empty:
    st.success(
        "No se detectaron contradicciones con las reglas automáticas actuales."
    )
else:
    st.error(
        f"Se detectaron {len(contradictions_df)} contradicción(es) potencial(es). "
        "Abre cada bloque para ver ambas versiones completas."
    )

    for index, row in contradictions_df.iterrows():
        with st.expander(
            f"🔴 Contradicción {index + 1}: {row.get('Tema', 'Tema no identificado')}",
            expanded=index == 0,
        ):
            left, right = st.columns(2)

            with left:
                st.markdown("### Versión 1 completa")
                st.write(row.get("Versión 1 completa", ""))
                st.markdown(
                    f"**Documento al que pertenece la versión 1:** "
                    f"{row.get('Documento de la versión 1', '')}"
                )
                st.markdown(
                    f"**Página:** {row.get('Página de la versión 1', '')}"
                )
                st.caption(
                    f"Método: {row.get('Método 1', '')} | "
                    f"OCR: {row.get('Confianza OCR 1', '')}"
                )

            with right:
                st.markdown("### Versión 2 completa")
                st.write(row.get("Versión 2 completa", ""))
                st.markdown(
                    f"**Documento al que pertenece la versión 2:** "
                    f"{row.get('Documento de la versión 2', '')}"
                )
                st.markdown(
                    f"**Página:** {row.get('Página de la versión 2', '')}"
                )
                st.caption(
                    f"Método: {row.get('Método 2', '')} | "
                    f"OCR: {row.get('Confianza OCR 2', '')}"
                )

            st.markdown(
                f"**Tipo de oposición:** {row.get('Tipo de oposición', '')}"
            )

            st.markdown("### Qué se contradice")
            st.write(row.get("Qué se contradice", ""))

            st.markdown("### Por qué importa")
            st.write(row.get("Por qué importa", ""))

            st.markdown("### Todo lo que debes revisar")
            st.warning(row.get("Todo lo que debe revisarse", ""))

            st.markdown("### Prueba que puede faltar")
            st.write(row.get("Prueba que puede faltar", ""))

            st.markdown("### Conclusión preliminar")
            st.info(row.get("Conclusión preliminar", ""))

    st.markdown("### Tabla editable de contradicciones")

    contradictions_df = st.data_editor(
        contradictions_df,
        use_container_width=True,
        hide_index=True,
        column_config={
            "Versión 1 completa": st.column_config.TextColumn(
                "Versión 1 completa",
                width="large",
            ),
            "Versión 2 completa": st.column_config.TextColumn(
                "Versión 2 completa",
                width="large",
            ),
            "Conclusión revisada": st.column_config.TextColumn(
                "Conclusión revisada",
                width="large",
            ),
            "Observaciones": st.column_config.TextColumn(
                "Observaciones",
                width="large",
            ),
        },
        disabled=[
            "Tema",
            "Versión 1 completa",
            "Fuente 1",
            "Método 1",
            "Confianza OCR 1",
            "Versión 2 completa",
            "Fuente 2",
            "Método 2",
            "Confianza OCR 2",
            "Tipo de oposición",
            "Evaluación",
        ],
        key="full_contradictions",
    )

st.subheader("4. Tabla de control")

table = pd.DataFrame(
    [
        {
            "Área": item.area,
            "Estado": badge(item.color),
            "Atención": item.score,
            "Qué debes cuidar": item.reason,
            "Acción sugerida": item.action,
            "Fuente": item.source,
            "Revisión humana": "",
        }
        for item in items
    ]
)

edited = st.data_editor(
    table,
    use_container_width=True,
    hide_index=True,
    column_config={
        "Atención": st.column_config.ProgressColumn(
            "Atención",
            min_value=0,
            max_value=100,
            format="%d",
        ),
        "Revisión humana": st.column_config.TextColumn(
            "Revisión humana",
            width="large",
        ),
    },
    key="semaphore_table",
)

st.subheader("5. Exportar reporte")

excel = io.BytesIO()
with pd.ExcelWriter(excel, engine="openpyxl") as writer:
    edited.to_excel(writer, sheet_name="Semaforo expediente", index=False)
    pd.DataFrame(quality_rows).to_excel(
        writer,
        sheet_name="Calidad documental",
        index=False,
    )
    contradictions_df.to_excel(
        writer,
        sheet_name="Contradicciones completas",
        index=False,
    )

download1, download2 = st.columns(2)

with download1:
    st.download_button(
        "Descargar semáforo y contradicciones en Excel",
        data=excel.getvalue(),
        file_name="semaforo_expediente_contradicciones.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True,
    )

with download2:
    st.download_button(
        "Descargar informe completo en Word",
        data=create_word(items, overall, contradictions_df),
        file_name="informe_semaforo_contradicciones.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )

st.warning(
    "Una contradicción automática es una señal de revisión, no una conclusión jurídica. "
    "Debe verificarse el contexto completo de ambas fuentes."
)


