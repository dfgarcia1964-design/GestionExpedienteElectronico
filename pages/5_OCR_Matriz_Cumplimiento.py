import io
import re
import unicodedata
from difflib import SequenceMatcher
from typing import Optional

import pandas as pd
import pytesseract
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from pdf2image import convert_from_bytes
from pdf_compat import PdfReader


st.set_page_config(
    page_title="OCR y matriz de cumplimiento",
    page_icon="🔎",
    layout="wide",
)

st.title("🔎 OCR y matriz de cumplimiento del fallo")
st.caption(
    "Lee documentos digitales o escaneados, extrae las órdenes del fallo "
    "y construye una matriz editable de cumplimiento."
)


PALABRAS_VACIAS = {
    "para", "como", "esta", "este", "estos", "estas", "desde", "hasta",
    "sobre", "entre", "dentro", "fuera", "ante", "bajo", "contra",
    "segun", "mediante", "porque", "cuando", "donde", "quien", "cual",
    "del", "las", "los", "una", "unos", "unas", "por", "con", "sin",
    "que", "sus", "son", "sea", "ser", "fue", "han", "hay", "más",
    "mas", "al", "se", "de", "la", "el", "en", "y", "o", "a",
}

VERBOS_ORDEN = (
    "ordenar",
    "ordena",
    "ordénese",
    "ordenese",
    "requerir",
    "requiere",
    "requiérase",
    "requierase",
    "disponer",
    "dispone",
    "autorizar",
    "autorice",
    "entregar",
    "garantizar",
    "realizar",
    "programar",
    "suministrar",
    "responder",
    "resolver",
    "remitir",
    "abstenerse",
    "vincular",
    "notificar",
)

EXPRESIONES_CUMPLIMIENTO = (
    "se dio cumplimiento",
    "dimos cumplimiento",
    "cumplimiento del fallo",
    "se cumplió",
    "se cumplio",
    "fue entregado",
    "se realizó",
    "se realizo",
    "se autorizó",
    "se autorizo",
    "se programó",
    "se programo",
    "se remitió",
    "se remitio",
    "se respondió",
    "se respondio",
    "se adjunta",
    "se aporta",
    "acta de entrega",
    "constancia de entrega",
)

EXPRESIONES_INCUMPLIMIENTO = (
    "no se ha cumplido",
    "incumplimiento",
    "no fue entregado",
    "no se realizó",
    "no se realizo",
    "no se autorizó",
    "no se autorizo",
    "no ha sido posible",
    "pendiente",
    "sin respuesta",
    "no existe prueba",
)


def normalizar(texto: str) -> str:
    texto = texto or ""
    texto = unicodedata.normalize("NFD", texto)
    texto = "".join(
        caracter
        for caracter in texto
        if unicodedata.category(caracter) != "Mn"
    )
    texto = texto.lower()
    texto = re.sub(r"\s+", " ", texto)
    return texto.strip()


def limpiar_texto(texto: str) -> str:
    texto = texto or ""
    texto = texto.replace("\x00", " ")
    texto = re.sub(r"[ \t]+", " ", texto)
    texto = re.sub(r"\n{3,}", "\n\n", texto)
    return texto.strip()


def contar_caracteres_utiles(texto: str) -> int:
    return len(re.sub(r"[^a-zA-ZáéíóúÁÉÍÓÚñÑ0-9]", "", texto or ""))


def ocr_pagina_pdf(
    contenido_pdf: bytes,
    numero_pagina: int,
    dpi: int,
) -> str:
    imagenes = convert_from_bytes(
        contenido_pdf,
        dpi=dpi,
        first_page=numero_pagina,
        last_page=numero_pagina,
        grayscale=True,
        thread_count=1,
        fmt="jpeg",
    )

    if not imagenes:
        return ""

    texto = pytesseract.image_to_string(
        imagenes[0],
        lang="spa",
        config="--oem 3 --psm 6",
    )

    return limpiar_texto(texto)


def extraer_pdf_con_ocr(
    contenido_pdf: bytes,
    usar_ocr: bool,
    minimo_caracteres: int,
    maximo_paginas_ocr: int,
    dpi: int,
) -> tuple[list[str], list[int], list[str]]:
    lector = PdfReader(io.BytesIO(contenido_pdf))
    paginas = []
    paginas_ocr = []
    advertencias = []

    for numero, pagina in enumerate(lector.pages, start=1):
        try:
            texto_digital = limpiar_texto(
                pagina.extract_text() or ""
            )
        except Exception:
            texto_digital = ""

        requiere_ocr = (
            usar_ocr
            and contar_caracteres_utiles(texto_digital)
            < minimo_caracteres
        )

        if requiere_ocr and numero <= maximo_paginas_ocr:
            try:
                texto_ocr = ocr_pagina_pdf(
                    contenido_pdf,
                    numero,
                    dpi,
                )

                if contar_caracteres_utiles(texto_ocr) > (
                    contar_caracteres_utiles(texto_digital)
                ):
                    texto_digital = texto_ocr
                    paginas_ocr.append(numero)

            except Exception as error:
                advertencias.append(
                    f"Página {numero}: OCR no disponible: {error}"
                )

        elif requiere_ocr and numero > maximo_paginas_ocr:
            advertencias.append(
                f"Página {numero}: no se aplicó OCR por el límite configurado."
            )

        paginas.append(texto_digital)

    return paginas, paginas_ocr, advertencias


def parece_fallo(texto: str) -> bool:
    contenido = normalizar(texto)

    expresiones = (
        "fallo de tutela",
        "administrando justicia",
        "en merito de lo expuesto",
        "resuelve",
        "amparar",
        "negar el amparo",
    )

    return sum(
        expresion in contenido
        for expresion in expresiones
    ) >= 2


def dividir_fragmentos(texto: str) -> list[str]:
    texto = limpiar_texto(texto)

    fragmentos = re.split(
        r"(?:\n\s*\n)|"
        r"(?<=[.;:])\s+(?=(?:primero|segundo|tercero|cuarto|quinto|"
        r"sexto|septimo|séptimo|octavo|noveno|decimo|décimo|ordenar|"
        r"ordenese|ordénese|requerir|requierase|requiérase)\b)",
        texto,
        flags=re.IGNORECASE,
    )

    return [
        re.sub(r"\s+", " ", fragmento).strip()
        for fragmento in fragmentos
        if len(fragmento.strip()) >= 25
    ]


def extraer_ordenes(
    paginas: list[str],
) -> list[dict]:
    ordenes = []
    numero_orden = 1
    inicio_resuelve = False

    for numero_pagina, pagina in enumerate(paginas, start=1):
        pagina_normalizada = normalizar(pagina)

        if "resuelve" in pagina_normalizada:
            inicio_resuelve = True

        if not inicio_resuelve:
            continue

        for fragmento in dividir_fragmentos(pagina):
            fragmento_normalizado = normalizar(fragmento)

            contiene_verbo = any(
                normalizar(verbo) in fragmento_normalizado
                for verbo in VERBOS_ORDEN
            )

            if not contiene_verbo:
                continue

            if len(fragmento) > 1800:
                fragmento = fragmento[:1800] + "..."

            ordenes.append(
                {
                    "N.º": numero_orden,
                    "Orden judicial": fragmento,
                    "Página del fallo": numero_pagina,
                }
            )

            numero_orden += 1

    if ordenes:
        return ordenes

    # Método alternativo cuando no se detecta claramente RESUELVE
    for numero_pagina, pagina in enumerate(paginas, start=1):
        for fragmento in dividir_fragmentos(pagina):
            fragmento_normalizado = normalizar(fragmento)

            if any(
                normalizar(verbo) in fragmento_normalizado
                for verbo in VERBOS_ORDEN
            ):
                ordenes.append(
                    {
                        "N.º": numero_orden,
                        "Orden judicial": fragmento[:1800],
                        "Página del fallo": numero_pagina,
                    }
                )
                numero_orden += 1

    return ordenes[:30]


def extraer_responsable(orden: str) -> str:
    patrones = [
        r"ordenar\s+a\s+(.{3,120}?)(?:\s+que\s+|\s+para\s+|,|\.)",
        r"ordenese\s+a\s+(.{3,120}?)(?:\s+que\s+|\s+para\s+|,|\.)",
        r"ordénese\s+a\s+(.{3,120}?)(?:\s+que\s+|\s+para\s+|,|\.)",
        r"requerir\s+a\s+(.{3,120}?)(?:\s+para\s+|\s+que\s+|,|\.)",
        r"requierase\s+a\s+(.{3,120}?)(?:\s+para\s+|\s+que\s+|,|\.)",
    ]

    for patron in patrones:
        coincidencia = re.search(
            patron,
            orden,
            flags=re.IGNORECASE,
        )

        if coincidencia:
            return re.sub(
                r"\s+",
                " ",
                coincidencia.group(1),
            ).strip(" ,.;:")

    return "Requiere identificación manual"


def extraer_plazo(orden: str) -> str:
    patrones = [
        r"(?:dentro de|en el termino de|en el término de|plazo de)\s+"
        r"(?:las\s+|los\s+)?(?:\d+|[a-záéíóúñ]+)\s+"
        r"(?:horas|dias|días)(?:\s+habiles|\s+hábiles)?",
        r"termino improrrogable de\s+"
        r"(?:\d+|[a-záéíóúñ]+)\s+"
        r"(?:horas|dias|días)(?:\s+habiles|\s+hábiles)?",
        r"de manera inmediata",
        r"inmediatamente",
    ]

    for patron in patrones:
        coincidencia = re.search(
            patron,
            orden,
            flags=re.IGNORECASE,
        )

        if coincidencia:
            return coincidencia.group(0)

    return "No detectado"


def palabras_significativas(texto: str) -> set[str]:
    palabras = re.findall(
        r"\b[a-záéíóúñ]{4,}\b",
        normalizar(texto),
    )

    return {
        palabra
        for palabra in palabras
        if palabra not in PALABRAS_VACIAS
    }


def similitud_textual(texto_a: str, texto_b: str) -> float:
    palabras_a = palabras_significativas(texto_a)
    palabras_b = palabras_significativas(texto_b)

    if not palabras_a or not palabras_b:
        return 0.0

    interseccion = palabras_a.intersection(palabras_b)
    union = palabras_a.union(palabras_b)

    similitud_jaccard = len(interseccion) / max(len(union), 1)

    secuencia_a = " ".join(sorted(palabras_a))
    secuencia_b = " ".join(sorted(palabras_b))

    similitud_secuencia = SequenceMatcher(
        None,
        secuencia_a,
        secuencia_b,
    ).ratio()

    return round(
        (similitud_jaccard * 0.70)
        + (similitud_secuencia * 0.30),
        4,
    )


def buscar_mejor_evidencia(
    orden: str,
    documentos: list[dict],
) -> dict:
    mejor = {
        "puntaje": 0.0,
        "documento": "",
        "pagina": "",
        "fragmento": "",
        "estado": "Sin prueba localizada",
    }

    for documento in documentos:
        for numero_pagina, pagina in enumerate(
            documento["paginas"],
            start=1,
        ):
            for fragmento in dividir_fragmentos(pagina):
                puntaje = similitud_textual(
                    orden,
                    fragmento,
                )

                if puntaje <= mejor["puntaje"]:
                    continue

                contenido = normalizar(fragmento)

                tiene_cumplimiento = any(
                    normalizar(expresion) in contenido
                    for expresion in EXPRESIONES_CUMPLIMIENTO
                )

                tiene_incumplimiento = any(
                    normalizar(expresion) in contenido
                    for expresion in EXPRESIONES_INCUMPLIMIENTO
                )

                if tiene_incumplimiento and puntaje >= 0.08:
                    estado = "Posible incumplimiento"
                elif tiene_cumplimiento and puntaje >= 0.08:
                    estado = "Posible cumplimiento"
                elif puntaje >= 0.12:
                    estado = "Requiere revisión"
                else:
                    estado = "Sin prueba localizada"

                mejor = {
                    "puntaje": puntaje,
                    "documento": documento["nombre"],
                    "pagina": numero_pagina,
                    "fragmento": fragmento[:900],
                    "estado": estado,
                }

    return mejor


def generar_excel(
    matriz: pd.DataFrame,
    resumen_documentos: pd.DataFrame,
) -> bytes:
    salida = io.BytesIO()

    with pd.ExcelWriter(
        salida,
        engine="openpyxl",
    ) as escritor:
        matriz.to_excel(
            escritor,
            sheet_name="Matriz cumplimiento",
            index=False,
        )

        resumen_documentos.to_excel(
            escritor,
            sheet_name="Documentos OCR",
            index=False,
        )

    return salida.getvalue()


def generar_word(
    matriz: pd.DataFrame,
) -> bytes:
    documento = Document()

    documento.styles["Normal"].font.name = "Arial"
    documento.styles["Normal"].font.size = Pt(10)

    titulo = documento.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    texto_titulo = titulo.add_run(
        "MATRIZ PRELIMINAR DE CUMPLIMIENTO DEL FALLO"
    )
    texto_titulo.bold = True
    texto_titulo.font.size = Pt(14)

    documento.add_paragraph(
        "Documento generado automáticamente. Cada conclusión debe "
        "contrastarse con el expediente original y revisarse jurídicamente."
    )

    for _, fila in matriz.iterrows():
        documento.add_heading(
            f"Orden {fila.get('N.º', '')}",
            level=2,
        )

        campos = [
            ("Orden judicial", fila.get("Orden judicial", "")),
            ("Página del fallo", fila.get("Página del fallo", "")),
            ("Responsable", fila.get("Responsable", "")),
            ("Plazo", fila.get("Plazo", "")),
            ("Evidencia localizada", fila.get("Evidencia localizada", "")),
            ("Página de evidencia", fila.get("Página de evidencia", "")),
            ("Estado", fila.get("Estado", "")),
            ("Nivel de coincidencia", fila.get("Coincidencia", "")),
            ("Observaciones", fila.get("Observaciones", "")),
        ]

        for etiqueta, valor in campos:
            parrafo = documento.add_paragraph()
            encabezado = parrafo.add_run(f"{etiqueta}: ")
            encabezado.bold = True
            parrafo.add_run(str(valor or ""))

    salida = io.BytesIO()
    documento.save(salida)
    return salida.getvalue()


with st.sidebar:
    st.header("Configuración OCR")

    usar_ocr = st.checkbox(
        "Aplicar OCR cuando el PDF no tenga texto",
        value=True,
    )

    minimo_caracteres = st.slider(
        "Mínimo de caracteres por página",
        min_value=20,
        max_value=300,
        value=80,
        step=10,
    )

    maximo_paginas_ocr = st.slider(
        "Máximo de páginas OCR por documento",
        min_value=5,
        max_value=100,
        value=30,
        step=5,
    )

    dpi = st.select_slider(
        "Calidad OCR",
        options=[150, 200, 250, 300],
        value=200,
    )

    st.warning(
        "Un DPI mayor puede mejorar la lectura, pero aumenta el tiempo "
        "y el consumo de memoria."
    )


archivos = st.file_uploader(
    "Sube el fallo y las piezas posteriores de cumplimiento",
    type=["pdf"],
    accept_multiple_files=True,
)

if not archivos:
    st.info(
        "Carga primero el fallo y después las respuestas, constancias, "
        "requerimientos, conceptos, actas o incidentes de desacato."
    )
    st.stop()


documentos = []
progreso = st.progress(0)
estado_proceso = st.empty()

for indice, archivo in enumerate(archivos):
    estado_proceso.info(
        f"Procesando: {archivo.name}"
    )

    try:
        contenido = archivo.getvalue()

        paginas, paginas_ocr, advertencias = extraer_pdf_con_ocr(
            contenido,
            usar_ocr,
            minimo_caracteres,
            maximo_paginas_ocr,
            dpi,
        )

        texto_total = "\n\n".join(paginas)

        documentos.append(
            {
                "nombre": archivo.name,
                "paginas": paginas,
                "texto": texto_total,
                "paginas_ocr": paginas_ocr,
                "advertencias": advertencias,
                "parece_fallo": parece_fallo(texto_total),
            }
        )

    except Exception as error:
        documentos.append(
            {
                "nombre": archivo.name,
                "paginas": [],
                "texto": "",
                "paginas_ocr": [],
                "advertencias": [str(error)],
                "parece_fallo": False,
            }
        )

    progreso.progress((indice + 1) / len(archivos))

estado_proceso.empty()


resumen_documentos = pd.DataFrame(
    [
        {
            "Documento": documento["nombre"],
            "Páginas": len(documento["paginas"]),
            "Páginas procesadas con OCR": (
                ", ".join(
                    str(numero)
                    for numero in documento["paginas_ocr"]
                )
                or "Ninguna"
            ),
            "Caracteres extraídos": len(documento["texto"]),
            "Posible fallo": (
                "Sí"
                if documento["parece_fallo"]
                else "No"
            ),
            "Advertencias": " | ".join(
                documento["advertencias"]
            ),
        }
        for documento in documentos
    ]
)

st.subheader("1. Resultado de lectura y OCR")

st.dataframe(
    resumen_documentos,
    use_container_width=True,
    hide_index=True,
)


nombres_documentos = [
    documento["nombre"]
    for documento in documentos
]

fallos_detectados = [
    documento["nombre"]
    for documento in documentos
    if documento["parece_fallo"]
]

indice_predeterminado = 0

if fallos_detectados:
    indice_predeterminado = nombres_documentos.index(
        fallos_detectados[0]
    )

fallo_seleccionado = st.selectbox(
    "Selecciona cuál documento contiene el fallo",
    options=nombres_documentos,
    index=indice_predeterminado,
)

documento_fallo = next(
    documento
    for documento in documentos
    if documento["nombre"] == fallo_seleccionado
)

documentos_evidencia = [
    documento
    for documento in documentos
    if documento["nombre"] != fallo_seleccionado
]


st.subheader("2. Órdenes detectadas")

ordenes = extraer_ordenes(
    documento_fallo["paginas"]
)

if not ordenes:
    st.error(
        "No se detectaron órdenes automáticamente. Verifica que el PDF "
        "correcto esté seleccionado y que el texto u OCR sea legible."
    )
    st.stop()

st.success(
    f"Se detectaron {len(ordenes)} posibles órdenes."
)


filas_matriz = []

for orden in ordenes:
    evidencia = buscar_mejor_evidencia(
        orden["Orden judicial"],
        documentos_evidencia,
    )

    filas_matriz.append(
        {
            "N.º": orden["N.º"],
            "Orden judicial": orden["Orden judicial"],
            "Página del fallo": orden["Página del fallo"],
            "Responsable": extraer_responsable(
                orden["Orden judicial"]
            ),
            "Plazo": extraer_plazo(
                orden["Orden judicial"]
            ),
            "Evidencia localizada": (
                f"{evidencia['documento']}: "
                f"{evidencia['fragmento']}"
                if evidencia["documento"]
                else ""
            ),
            "Página de evidencia": evidencia["pagina"],
            "Estado": evidencia["estado"],
            "Coincidencia": (
                f"{evidencia['puntaje'] * 100:.1f}%"
            ),
            "Observaciones": "",
        }
    )

matriz_inicial = pd.DataFrame(filas_matriz)


st.subheader("3. Matriz editable de cumplimiento")

st.info(
    "Revisa cada fila. Puedes corregir responsable, plazo, estado, "
    "evidencia y observaciones antes de descargar el informe."
)

estados_permitidos = [
    "Cumplida",
    "Parcialmente cumplida",
    "Incumplida",
    "No verificable",
    "Requiere revisión",
    "Posible cumplimiento",
    "Posible incumplimiento",
    "Sin prueba localizada",
]

matriz_editada = st.data_editor(
    matriz_inicial,
    use_container_width=True,
    hide_index=True,
    num_rows="dynamic",
    column_config={
        "N.º": st.column_config.NumberColumn(
            "N.º",
            disabled=True,
        ),
        "Página del fallo": st.column_config.NumberColumn(
            "Página del fallo",
            min_value=1,
        ),
        "Página de evidencia": st.column_config.NumberColumn(
            "Página de evidencia",
            min_value=1,
        ),
        "Estado": st.column_config.SelectboxColumn(
            "Estado",
            options=estados_permitidos,
            required=True,
        ),
        "Orden judicial": st.column_config.TextColumn(
            "Orden judicial",
            width="large",
        ),
        "Evidencia localizada": st.column_config.TextColumn(
            "Evidencia localizada",
            width="large",
        ),
        "Observaciones": st.column_config.TextColumn(
            "Observaciones",
            width="large",
        ),
    },
    key="matriz_cumplimiento",
)


st.subheader("4. Resumen de cumplimiento")

conteo_estados = (
    matriz_editada["Estado"]
    .fillna("Sin estado")
    .value_counts()
    .rename_axis("Estado")
    .reset_index(name="Cantidad")
)

st.dataframe(
    conteo_estados,
    use_container_width=True,
    hide_index=True,
)


columna_excel, columna_word = st.columns(2)

with columna_excel:
    st.download_button(
        "Descargar matriz en Excel",
        data=generar_excel(
            matriz_editada,
            resumen_documentos,
        ),
        file_name="matriz_cumplimiento_fallo.xlsx",
        mime=(
            "application/vnd.openxmlformats-officedocument."
            "spreadsheetml.sheet"
        ),
        use_container_width=True,
    )

with columna_word:
    st.download_button(
        "Descargar matriz en Word",
        data=generar_word(matriz_editada),
        file_name="matriz_cumplimiento_fallo.docx",
        mime=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
        use_container_width=True,
    )


st.warning(
    "La coincidencia automática no demuestra cumplimiento jurídico. "
    "Una prueba puede referirse al mismo tema sin acreditar que la orden "
    "fue cumplida integralmente, dentro del plazo y por el responsable."
)
